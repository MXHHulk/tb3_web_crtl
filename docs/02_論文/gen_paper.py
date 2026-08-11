# -*- coding: utf-8 -*-
"""
依本專案內容生成兩份國際期刊投稿格式論文（中文版 / 英文版）Word 檔。
輸出：
  docs/02_論文/paper_CN.docx
  docs/02_論文/paper_EN.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.realpath(__file__))

# ────────────────────────────────────────────────────────────
#  共用樣式工具
# ────────────────────────────────────────────────────────────
def set_cjk(run, font_cn="標楷體", font_en="Times New Roman"):
    """同時設定西文與中日韓字型。"""
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)


def base_doc(cjk=False):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    if cjk:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), "新細明體")
    # 邊界
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
    return doc


def add_title(doc, text, cjk=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17)
    if cjk:
        set_cjk(r, "標楷體", "Times New Roman")
    p.paragraph_format.space_after = Pt(6)
    return p


def add_center(doc, text, size=10.5, italic=False, bold=False, cjk=False, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if cjk:
        set_cjk(r)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text, cjk=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if cjk:
        set_cjk(r, "標楷體", "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_sub(doc, text, cjk=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10.5)
    if cjk:
        set_cjk(r, "標楷體", "Times New Roman")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text, cjk=False, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    if cjk:
        set_cjk(r)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_abstract(doc, label, text, cjk=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rl = p.add_run(label)
    rl.bold = True
    rl.font.size = Pt(10)
    rt = p.add_run(text)
    rt.font.size = Pt(10)
    if cjk:
        set_cjk(rl); set_cjk(rt)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    return p


def add_hyperlink(paragraph, url, text, size=9):
    """在段落尾端加入可點擊的外部超連結（藍色底線）。"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_ref(doc, idx, text, url, label="Available: ", cjk=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{idx}] {text} {label}")
    r.font.size = Pt(9)
    if cjk:
        set_cjk(r)
    add_hyperlink(p, url, url, size=9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    return p


def make_two_column(doc):
    """正文改為雙欄（期刊常見排版）。"""
    sect = doc.add_section(WD_SECTION.CONTINUOUS)
    cols = sect._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')
    return sect


REFS = [
    ('E. Galceran and M. Carreras, "A survey on coverage path planning for robotics," Robotics and Autonomous Systems, vol. 61, no. 12, pp. 1258–1276, 2013.',
     'https://www.semanticscholar.org/paper/a9c7f3a32865088a1113ba415178cbd14db2b663'),
    ('H. Choset and P. Pignon, "Coverage path planning: The boustrophedon cellular decomposition," in Field and Service Robotics, Springer, 1998, pp. 203–209.',
     'https://publications.ri.cmu.edu/storage/publications/pub_files/pub1/choset_howie_1997_5/choset_howie_1997_5.pdf'),
    ('G. Grisetti, C. Stachniss, and W. Burgard, "Improved techniques for grid mapping with Rao-Blackwellized particle filters," IEEE Transactions on Robotics, vol. 23, no. 1, pp. 34–46, 2007.',
     'https://www.ipb.uni-bonn.de/wp-content/papercite-data/pdf/grisetti07tro.pdf'),
    ('D. Fox, W. Burgard, and S. Thrun, "The dynamic window approach to collision avoidance," IEEE Robotics & Automation Magazine, vol. 4, no. 1, pp. 23–33, 1997.',
     'https://www.ri.cmu.edu/publications/the-dynamic-window-approach-to-collision-avoidance/'),
    ('M. Quigley et al., "ROS: an open-source Robot Operating System," in ICRA Workshop on Open Source Software, 2009.',
     'http://ai.stanford.edu/~mquigley/papers/icra2009-ros.pdf'),
    ('R. Bormann, F. Jordan, J. Hampp, and M. Hagele, "Indoor coverage path planning: Survey, implementation, analysis," in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2018, pp. 1718–1725.',
     'https://publica.fraunhofer.de/entities/publication/f537c15d-4cbe-4672-9d86-6e756a9ce71b'),
    ('I. T. Jolliffe and J. Cadima, "Principal component analysis: a review and recent developments," Philosophical Transactions of the Royal Society A, vol. 374, no. 2065, 2016.',
     'https://doi.org/10.1098/rsta.2015.0202'),
]

# ════════════════════════════════════════════════════════════
#  英文版
# ════════════════════════════════════════════════════════════
def build_en():
    doc = base_doc(cjk=False)
    add_title(doc, "A ROS-Based Complete Coverage Path Planning System for "
                   "TurtleBot3 Using PCA-Aligned Boustrophedon Decomposition "
                   "with Real-Time Web Visualization")
    add_center(doc, "MXHHulk", size=11, bold=True)
    add_center(doc, "Department of Robotics and Automation", size=10, italic=True)
    add_center(doc, "hulk71777@gmail.com", size=10, italic=True, after=10)

    add_abstract(doc, "Abstract—",
        "Complete coverage path planning (CCPP) requires a mobile robot to "
        "traverse every reachable region of a workspace, a capability central "
        "to cleaning robots, inspection, and agricultural automation. This paper "
        "presents a lightweight, modular CCPP system built on the Robot Operating "
        "System (ROS Noetic) for the TurtleBot3 platform. The system couples "
        "gmapping-based simultaneous localization and mapping (SLAM) with a "
        "boustrophedon (back-and-forth) coverage strategy whose sweep direction "
        "is automatically aligned to the dominant geometric axis of the free space "
        "through principal component analysis (PCA). This alignment eliminates the "
        "staircase artifacts and excessive turning that arise when sweep lines are "
        "fixed to the map axes while walls are oblique. A morphological safety "
        "margin keeps the planned path clear of obstacles, and the resulting "
        "waypoints are executed through the move_base navigation stack with "
        "Dynamic Window Approach (DWA) local obstacle avoidance and per-waypoint "
        "fault tolerance. A core design principle is the separation of the planning "
        "algorithm from ROS, yielding a pure, independently testable module shared "
        "by multiple nodes. A Flask-based web interface provides real-time, "
        "multi-layer visualization of the map, robot trajectory, and coverage path "
        "with one-click task control. Experiments on a physical TurtleBot3 Burger "
        "confirm that the system produces parallel, gap-free coverage paths that "
        "adapt to arbitrarily oriented rooms.")
    add_abstract(doc, "Keywords—",
        "complete coverage path planning; boustrophedon; principal component "
        "analysis; ROS; SLAM; mobile robot; TurtleBot3; web visualization")

    make_two_column(doc)

    add_heading(doc, "I.  INTRODUCTION")
    add_body(doc,
        "Unlike point-to-point navigation, which seeks a single path from a start "
        "to a goal, complete coverage path planning (CCPP) demands that the robot "
        "pass over the entire free area of its workspace, ideally without omission "
        "or redundant revisiting [1]. The most familiar instance is the "
        "household cleaning robot, but the same problem underlies floor inspection, "
        "lawn mowing, painting, and agricultural field operations.")
    add_body(doc,
        "Among coverage strategies, the boustrophedon (“as the ox plows”) "
        "pattern is widely adopted for its simplicity and near-optimal behavior in "
        "convex sub-regions: the robot sweeps a straight line, turns, and sweeps a "
        "parallel line in the opposite direction [2]. A practical difficulty arises "
        "when the strategy is applied directly to a map produced by SLAM: because "
        "the map frame is fixed by the robot's heading at start-up, room walls are "
        "frequently oblique to the map axes. Sweeping along a fixed axis then yields "
        "ragged, staircase-shaped boundaries and an excessive number of turns, "
        "degrading both efficiency and coverage quality.")
    add_body(doc,
        "This paper describes a complete, deployable CCPP system that addresses "
        "this issue and integrates the full pipeline from SLAM to execution and "
        "visualization. The main contributions are: (i) a PCA-based automatic "
        "alignment of the sweep direction to the principal axis of the free space, "
        "making coverage robust to room orientation; (ii) a clean separation of the "
        "coverage algorithm from ROS, producing a dependency-free, independently "
        "testable module reused across nodes; and (iii) a real-time, multi-layer "
        "web interface for monitoring and control. The system runs on commodity "
        "TurtleBot3 hardware using only open-source components.")

    add_heading(doc, "II.  RELATED WORK")
    add_body(doc,
        "Coverage path planning has been surveyed extensively [1], [6]. "
        "Exact cellular decomposition methods, of which the boustrophedon "
        "decomposition is the canonical example [2], partition the free space into "
        "cells that admit simple back-and-forth coverage. Our work adopts the "
        "boustrophedon sweep but, rather than performing a full critical-point "
        "decomposition, applies principal component analysis [7] to orient the "
        "sweep with the dominant direction of the reachable region, a pragmatic "
        "choice well suited to single connected rooms.")
    add_body(doc,
        "On the systems side, the Robot Operating System [5] provides the "
        "publish/subscribe and action infrastructure used here, gmapping [3] "
        "supplies the occupancy grid through Rao-Blackwellized particle-filter "
        "SLAM, and the move_base navigation stack with the Dynamic Window Approach "
        "[4] performs local collision avoidance during execution. Compared with "
        "general coverage frameworks [6], the present system is intentionally "
        "minimal and education-oriented, emphasizing a transparent, modular "
        "implementation and live visual feedback.")

    add_heading(doc, "III.  SYSTEM ARCHITECTURE")
    add_body(doc,
        "The system is organized as a small set of cooperating ROS nodes. A "
        "physical TurtleBot3 publishes laser scans (/scan) and odometry (/odom). "
        "The gmapping node consumes the scans and produces an occupancy grid "
        "(/map) together with the map→odom coordinate transform. Two custom "
        "nodes subscribe to this data. Module B, the planning node, computes the "
        "coverage path and republishes it as a nav_msgs/Path on /coverage_path for "
        "visualization in RViz. Module A, the map server, converts the grid into "
        "web image layers, records the robot trajectory, hosts the web interface, "
        "and dispatches the coverage waypoints to move_base for execution.")
    add_body(doc,
        "Both modules share a single, ROS-independent algorithm module "
        "(coverage_planner), which contains the entire planning logic in terms of "
        "plain NumPy arrays and world coordinates. Because it has no ROS "
        "dependency, it can be unit-tested in isolation and reused by any node. "
        "Execution goals are delegated to move_base through the ROS action "
        "interface, which is appropriate for long-running, cancellable tasks, "
        "whereas continuous data streams use the publish/subscribe interface.")

    add_heading(doc, "IV.  METHODOLOGY")
    add_sub(doc, "A.  Free-Space Extraction and Safety Margin")
    add_body(doc,
        "The occupancy grid encodes each cell as free (0), occupied (100), or "
        "unknown (-1). Before planning, obstacles are dilated by a safety margin "
        "of approximately one robot radius (0.15 m) using a morphological "
        "structuring element whose radius in cells is the margin divided by the "
        "map resolution. The traversable set is then defined as the cells that are "
        "free and not within the dilated obstacle region, ensuring the planned "
        "path never grazes a wall.")
    add_sub(doc, "B.  PCA-Aligned Sweep Direction")
    add_body(doc,
        "Let the world coordinates of all traversable cells form a point set. The "
        "covariance matrix of these points is computed and its eigenvectors are "
        "obtained. The eigenvector associated with the largest eigenvalue defines "
        "the long axis of the free space and is used as the sweep direction, "
        "minimizing the number of turns; the eigenvector of the smallest "
        "eigenvalue defines the step direction along which successive sweep lines "
        "advance. Because the axes are derived from the data, the sweep lines "
        "remain parallel to the walls regardless of how the map is rotated.")
    add_sub(doc, "C.  Boustrophedon Waypoint Generation")
    add_body(doc,
        "Each traversable point is projected onto the long and short axes. "
        "Starting half a spacing inside the boundary for symmetric coverage, sweep "
        "lines are placed at a fixed interval (spacing = 0.25 m, smaller than the "
        "robot diameter to guarantee overlap). For each line, the two extreme "
        "projections define the segment endpoints. A boolean flag alternates the "
        "traversal direction of consecutive lines, reproducing the characteristic "
        "back-and-forth boustrophedon motion. The output is an ordered list of "
        "world-coordinate waypoints.")
    add_sub(doc, "D.  Execution and Fault Tolerance")
    add_body(doc,
        "The map server sends waypoints to move_base one at a time, orienting each "
        "goal toward the next waypoint. move_base, configured with the DWA local "
        "planner, plans a global route and continuously samples feasible velocity "
        "commands within the dynamic window to follow it while avoiding obstacles. "
        "If a waypoint cannot be reached (the action returns a non-successful "
        "status), the system logs a warning and proceeds to the next waypoint "
        "instead of stalling, so a single unreachable goal does not abort the task.")

    add_heading(doc, "V.  IMPLEMENTATION")
    add_body(doc,
        "The map server simultaneously runs ROS callback threads, a Flask web "
        "server thread, and a background execution thread. Shared state—the map "
        "layers, the robot pose and trajectory, and the coverage task status—is "
        "protected by three independent locks to prevent data races. Task "
        "termination uses cooperative cancellation: the execution loop periodically "
        "inspects the task state and, upon a stop request, cancels outstanding "
        "goals and exits cleanly rather than being forcibly killed.")
    add_body(doc,
        "For visualization, the occupancy grid is rendered into three layers—the "
        "raw map, a morphologically eroded map, and a dilated map—and cropped to "
        "the bounding box of the known region to reduce transmission. The browser "
        "client overlays these layers on an HTML5 canvas together with the coverage "
        "path, the traveled trajectory, and the live robot position, polling the "
        "server once per second. The interface exposes one-click start and stop "
        "controls for the coverage task.")

    add_heading(doc, "VI.  RESULTS AND DISCUSSION")
    add_body(doc,
        "The system was deployed on a physical TurtleBot3 Burger running ROS "
        "Noetic. After teleoperating the robot to build a sufficiently complete "
        "map with gmapping, the planner generated coverage paths whose sweep lines "
        "were consistently parallel to the room walls, even when the map frame was "
        "rotated relative to the room. This confirms that PCA alignment removes the "
        "staircase artifacts that a fixed-axis sweep would produce, reducing the "
        "number of turns and improving coverage uniformity.")
    add_body(doc,
        "Qualitatively, the safety-margin dilation kept the executed trajectory "
        "clear of walls, while the per-waypoint fault tolerance allowed the robot "
        "to complete coverage even when isolated goals were temporarily "
        "unreachable due to localization drift or transient obstacles. The "
        "decoupling of the planning algorithm from ROS proved valuable in "
        "practice: the coverage logic could be exercised and tuned independently "
        "of the robot. Limitations include reliance on a single connected free "
        "region (the PCA alignment is most effective for a dominant rectangular "
        "area) and a fixed S-shaped waypoint ordering that is not globally "
        "optimized for travel distance.")

    add_heading(doc, "VII.  CONCLUSION")
    add_body(doc,
        "We presented a modular, ROS-based complete coverage path planning system "
        "for TurtleBot3 that combines PCA-aligned boustrophedon planning, a "
        "morphological safety margin, fault-tolerant execution through move_base, "
        "and a real-time multi-layer web interface. The PCA alignment makes "
        "coverage robust to room orientation, and the strict separation of "
        "algorithm from middleware yields a reusable, testable core. Future work "
        "includes multi-region decomposition for non-convex environments, "
        "travel-distance optimization of the waypoint order toward a traveling-"
        "salesman formulation, frontier-based autonomous exploration to remove the "
        "manual mapping step, and an online coverage-rate metric reported in the "
        "web interface.")

    add_heading(doc, "REFERENCES")
    for i, (text, url) in enumerate(REFS, 1):
        add_ref(doc, i, text, url, label="Available: ")

    out = os.path.join(HERE, "paper_EN.docx")
    doc.save(out)
    return out


# ════════════════════════════════════════════════════════════
#  中文版
# ════════════════════════════════════════════════════════════
def build_cn():
    doc = base_doc(cjk=True)
    add_title(doc, "基於 ROS 與主成分分析對齊牛耕式分解之 TurtleBot3 "
                   "完全覆蓋路徑規劃系統及其即時網頁視覺化", cjk=True)
    add_center(doc, "MXHHulk", size=11, bold=True, cjk=True)
    add_center(doc, "機器人與自動化研究", size=10, italic=True, cjk=True)
    add_center(doc, "hulk71777@gmail.com", size=10, italic=True, after=10)

    add_abstract(doc, "摘要—",
        "完全覆蓋路徑規劃（Complete Coverage Path Planning, CCPP）要求移動機器人"
        "走遍工作空間中每一個可到達的區域，是掃地機器人、巡檢與農業自動化的核心"
        "能力。本文提出一套以機器人作業系統（ROS Noetic）為基礎、適用於 TurtleBot3 "
        "平台的輕量化模組化 CCPP 系統。系統結合 gmapping 之同時定位與建圖（SLAM），"
        "並採用牛耕式（來回往復）覆蓋策略；其掃描方向透過主成分分析（PCA）自動對齊"
        "自由空間的主要幾何軸向。此一對齊機制可消除在牆面傾斜、掃描線卻固定於地圖"
        "軸向時所產生的階梯狀鋸齒與過量轉彎。系統並以形態學安全邊距使規劃路徑遠離"
        "障礙，所得路點再透過 move_base 導航堆疊執行，搭配動態視窗法（DWA）進行局部"
        "避障，並具備逐點失敗容錯。核心設計原則為「演算法與 ROS 解耦」，形成一個可"
        "獨立測試、可被多個節點共用的純演算法模組。系統另提供以 Flask 為基礎的網頁"
        "介面，可即時多圖層顯示地圖、機器人軌跡與覆蓋路徑，並支援一鍵任務控制。於"
        "實體 TurtleBot3 Burger 之實驗證實，本系統能產生平行且無遺漏的覆蓋路徑，並"
        "可自動適應任意朝向的房間。")
    add_abstract(doc, "關鍵詞—",
        "完全覆蓋路徑規劃；牛耕式；主成分分析；ROS；SLAM；移動機器人；TurtleBot3；"
        "網頁視覺化", cjk=True)

    make_two_column(doc)

    add_heading(doc, "一、緒論", cjk=True)
    add_body(doc,
        "與點對點導航（自起點求得抵達目標的單一路徑）不同，完全覆蓋路徑規劃（CCPP）"
        "要求機器人經過工作空間中所有自由區域，理想上不遺漏、不重複 [1]。最為"
        "人熟知的例子即家用掃地機器人，但相同問題亦存在於地面巡檢、割草、噴漆與農田"
        "作業之中。", cjk=True)
    add_body(doc,
        "在各種覆蓋策略中，牛耕式（取自「牛犁田」之意）因其簡潔且在凸子區域內近乎"
        "最佳的特性而被廣泛採用：機器人沿直線掃描、轉彎，再以相反方向掃描一條平行"
        "線 [2]。然而將其直接套用於 SLAM 所建之地圖時會遭遇實務困難：由於地圖座標"
        "系由機器人開機瞬間的朝向決定，房間牆面常與地圖軸向呈傾斜。此時沿固定軸向"
        "掃描將產生破碎的階梯狀邊界與大量轉彎，同時降低效率與覆蓋品質。", cjk=True)
    add_body(doc,
        "本文描述一套完整、可部署的 CCPP 系統，以解決上述問題，並整合自 SLAM 至"
        "執行與視覺化之完整流程。主要貢獻為：（一）以 PCA 將掃描方向自動對齊自由"
        "空間之主軸，使覆蓋對房間朝向具備強健性；（二）將覆蓋演算法與 ROS 徹底"
        "解耦，形成無相依、可獨立測試並跨節點重用之模組；（三）提供即時多圖層之"
        "網頁監控與控制介面。整套系統僅以開源元件運行於常見之 TurtleBot3 硬體。",
        cjk=True)

    add_heading(doc, "二、相關研究", cjk=True)
    add_body(doc,
        "覆蓋路徑規劃已有廣泛之文獻回顧 [1], [6]。精確胞元分解法（其中牛耕式"
        "分解為典型代表 [2]）將自由空間切分為可施以簡單來回覆蓋之胞元。本研究採用"
        "牛耕式掃描，但並未執行完整的臨界點分解，而是以主成分分析 [7] 將掃描方向"
        "對齊可達區域之主要方向；此一務實作法特別適用於單一連通房間。", cjk=True)
    add_body(doc,
        "於系統層面，機器人作業系統 [5] 提供本文所用之發布/訂閱與動作（action）"
        "基礎建設；gmapping [3] 以 Rao-Blackwellized 粒子濾波 SLAM 產生佔據網格；"
        "move_base 導航堆疊搭配動態視窗法 [4] 於執行期間進行局部避障。相較於通用"
        "覆蓋框架 [6]，本系統刻意保持精簡且偏重教學，著重於透明、模組化之實作與"
        "即時視覺回饋。", cjk=True)

    add_heading(doc, "三、系統架構", cjk=True)
    add_body(doc,
        "系統由一組互相協作的 ROS 節點構成。實體 TurtleBot3 發布雷射掃描（/scan）"
        "與里程計（/odom）。gmapping 節點消耗掃描資料，產生佔據網格（/map）及 "
        "map→odom 座標轉換。兩個自製節點訂閱此資料。模組 B（規劃節點）計算覆蓋"
        "路徑，並以 nav_msgs/Path 重新發布於 /coverage_path，供 RViz 顯示。模組 A"
        "（地圖伺服器）將網格轉為網頁圖層、記錄機器人軌跡、託管網頁介面，並將覆蓋"
        "路點派送至 move_base 執行。", cjk=True)
    add_body(doc,
        "兩個模組共用單一、與 ROS 無關之演算法模組（coverage_planner），其以普通 "
        "NumPy 陣列與世界座標承載全部規劃邏輯。由於不相依於 ROS，該模組可被獨立"
        "單元測試並為任一節點重用。執行目標透過 ROS 動作介面委派給 move_base，該"
        "介面適用於長時間、可取消之任務；而連續資料流則使用發布/訂閱介面。",
        cjk=True)

    add_heading(doc, "四、方法", cjk=True)
    add_sub(doc, "（甲）自由空間擷取與安全邊距", cjk=True)
    add_body(doc,
        "佔據網格將每格標示為自由（0）、佔據（100）或未知（-1）。規劃前，先以"
        "約一個機器人半徑（0.15 公尺）之安全邊距對障礙物進行膨脹，其結構元素之"
        "格半徑為邊距除以地圖解析度。可走集合定義為「自由且不落於膨脹後障礙區」"
        "之格子，以確保規劃路徑絕不貼牆。", cjk=True)
    add_sub(doc, "（乙）PCA 對齊之掃描方向", cjk=True)
    add_body(doc,
        "令所有可走格之世界座標構成一點集。計算其共變異數矩陣並求取特徵向量。"
        "對應最大特徵值之特徵向量定義自由空間之長軸，作為掃描方向以最小化轉彎"
        "次數；最小特徵值之特徵向量則定義短軸，供連續掃描線沿之前進。由於軸向"
        "源自資料本身，無論地圖如何旋轉，掃描線皆能保持與牆面平行。", cjk=True)
    add_sub(doc, "（丙）牛耕路點生成", cjk=True)
    add_body(doc,
        "將每個可走點投影至長軸與短軸。為使覆蓋對稱，第一條掃描線自邊界內縮半個"
        "間距，並以固定間距（spacing = 0.25 公尺，小於機器人直徑以確保重疊）佈設"
        "掃描線。每條線上以兩個極值投影定出線段端點。一個布林旗標使相鄰掃描線之"
        "走訪方向交替，重現牛耕式之來回特性。輸出為依序排列之世界座標路點清單。",
        cjk=True)
    add_sub(doc, "（丁）執行與失敗容錯", cjk=True)
    add_body(doc,
        "地圖伺服器逐一將路點送至 move_base，並使每個目標朝向下一路點。配置 DWA "
        "局部規劃器之 move_base 會規劃全域路線，並於動態視窗內持續取樣可行速度"
        "指令以沿線行進並避障。若某路點無法到達（動作回傳非成功狀態），系統記錄"
        "警告並前進至下一路點而非停滯，因此單一無法到達之目標不致中止整體任務。",
        cjk=True)

    add_heading(doc, "五、實作", cjk=True)
    add_body(doc,
        "地圖伺服器同時運行 ROS 回呼執行緒、Flask 網頁伺服器執行緒與背景執行"
        "執行緒。共享狀態—地圖圖層、機器人位姿與軌跡、覆蓋任務狀態—以三把"
        "獨立鎖保護以避免資料競爭。任務終止採用合作式取消：執行迴圈定期檢視"
        "任務狀態，於收到停止請求時取消未完成之目標並乾淨退出，而非被強制終止。",
        cjk=True)
    add_body(doc,
        "於視覺化方面，佔據網格被繪製為三種圖層—原始地圖、形態學侵蝕地圖與"
        "膨脹地圖—並裁切至已知區域之外接矩形以減少傳輸量。瀏覽器端於 HTML5 "
        "畫布上將上述圖層與覆蓋路徑、已行走軌跡及即時機器人位置疊加顯示，並"
        "每秒向伺服器輪詢一次。介面提供覆蓋任務之一鍵開始與停止控制。", cjk=True)

    add_heading(doc, "六、結果與討論", cjk=True)
    add_body(doc,
        "系統部署於運行 ROS Noetic 之實體 TurtleBot3 Burger。在以遙控方式驅動"
        "機器人並透過 gmapping 建立足夠完整之地圖後，規劃器所生成之覆蓋路徑，"
        "其掃描線均能穩定地與房間牆面平行，即使地圖座標系相對房間旋轉亦然。"
        "此結果證實 PCA 對齊可消除固定軸向掃描所造成之階梯鋸齒，減少轉彎次數"
        "並提升覆蓋均勻度。", cjk=True)
    add_body(doc,
        "就定性觀察而言，安全邊距膨脹使執行軌跡遠離牆面；逐點失敗容錯則使機器人"
        "即使在因定位漂移或暫時性障礙而導致個別目標一時無法到達時，仍能完成覆蓋。"
        "演算法與 ROS 之解耦於實務上極具價值：覆蓋邏輯得以獨立於機器人進行演練與"
        "調校。本系統之限制包括：仰賴單一連通自由區域（PCA 對齊對具主導性之矩形"
        "區域最為有效），以及固定之 S 形路點排序尚未就總行程進行全域最佳化。",
        cjk=True)

    add_heading(doc, "七、結論", cjk=True)
    add_body(doc,
        "本文提出一套模組化、以 ROS 為基礎之 TurtleBot3 完全覆蓋路徑規劃系統，"
        "結合 PCA 對齊之牛耕式規劃、形態學安全邊距、透過 move_base 之容錯執行，"
        "以及即時多圖層網頁介面。PCA 對齊使覆蓋對房間朝向具強健性，而演算法與"
        "中介軟體之嚴格分離則產生可重用、可測試之核心。未來工作包括：針對非凸"
        "環境之多區域分解、朝旅行推銷員問題形式對路點順序進行行程最佳化、以"
        "frontier 為基礎之自主探索以省去人工建圖步驟，以及於網頁介面回報之線上"
        "覆蓋率指標。", cjk=True)

    add_heading(doc, "參考文獻", cjk=True)
    for i, (text, url) in enumerate(REFS, 1):
        add_ref(doc, i, text, url, label="可取得： ", cjk=True)

    out = os.path.join(HERE, "paper_CN.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    en = build_en()
    cn = build_cn()
    print("生成完成：")
    print(" ", en)
    print(" ", cn)
