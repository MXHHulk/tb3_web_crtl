# -*- coding: utf-8 -*-
"""
依本專案內容生成一份「學習導向」的國際期刊論文格式 Word 檔（繁體中文）。

與 gen_paper.py（精簡投稿版）不同，本檔的目標是「教學」：
從零開始、由淺入深，讓完全不了解本專案的人也能讀懂整套系統，
同時保留國際期刊論文的版面結構（標題 / 摘要 / 關鍵詞 / 編號章節 / 參考文獻）。

所有參數均依 rebuild 分支「當前程式碼」撰寫：
  SPACING = 0.18 m, MARGIN = 0.10 m, CROP_PAD = 10, REPLAN_INTERVAL = 5.0 s。

輸出：docs/02_論文/學習論文_CN.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.realpath(__file__))

CN_FONT   = "標楷體"          # 中文標題/內文字型
CN_BODY   = "新細明體"        # 中文內文（明體較易閱讀長文）
EN_FONT   = "Times New Roman" # 西文字型
MONO_EN   = "Consolas"        # 程式碼西文
MONO_CN   = "新細明體"        # 程式碼/圖內中文與框線字元（等全形寬）


# ────────────────────────────────────────────────────────────
#  低階字型 / 樣式工具
# ────────────────────────────────────────────────────────────
def _set_fonts(run, cn, en):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)
    rfonts.set(qn('w:eastAsia'), cn)


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _box(paragraph, color="BBBBBB", sz=4):
    """為段落加上四周細框（用於程式碼/圖塊）。"""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '6')
        e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)


def base_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = EN_FONT
    st.font.size = Pt(11)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CN_BODY)
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)
    return doc


# ────────────────────────────────────────────────────────────
#  區塊層級 API
# ────────────────────────────────────────────────────────────
def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    _set_fonts(r, CN_FONT, EN_FONT)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1


def center(doc, text, size=11, italic=False, bold=False, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    _set_fonts(r, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(after)


def abstract(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rl = p.add_run(label)
    rl.bold = True
    rl.font.size = Pt(10.5)
    _set_fonts(rl, CN_FONT, EN_FONT)
    rt = p.add_run(text)
    rt.font.size = Pt(10.5)
    _set_fonts(rt, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    _set_fonts(r, CN_FONT, EN_FONT)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x5A, 0x88)
    _set_fonts(r, CN_FONT, EN_FONT)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(11)
    _set_fonts(r, CN_FONT, EN_FONT)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)
    _set_fonts(r, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.28)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    r = p.add_run(text)
    r.font.size = Pt(11)
    _set_fonts(r, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.size = Pt(11)
    _set_fonts(r, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    return p


def code(doc, text):
    """以等寬字 + 灰底 + 細框呈現的程式碼/虛擬碼/ASCII 圖塊。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.1)
    _shade(p, "F4F5F7")
    _box(p, "C9CDD3", 4)
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.size = Pt(9.5)
        _set_fonts(r, MONO_CN, MONO_EN)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_fonts(r, CN_BODY, EN_FONT)
    p.paragraph_format.space_after = Pt(10)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        para = hdr[j].paragraphs[0]
        run = para.add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        _set_fonts(run, CN_FONT, EN_FONT)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            para = cells[j].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            _set_fonts(run, CN_BODY, EN_FONT)
            para.paragraph_format.space_after = Pt(1)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_hyperlink(paragraph, url, text, size=9):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = OxmlElement('w:hyperlink')
    h.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; run.append(t)
    h.append(run)
    paragraph._p.append(h)


def ref(doc, idx, text, url):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{idx}] {text} 可取得： ")
    r.font.size = Pt(9)
    _set_fonts(r, CN_BODY, EN_FONT)
    add_hyperlink(p, url, url, size=9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)


def pagebreak(doc):
    doc.add_page_break()


REFS = [
    ('E. Galceran and M. Carreras, "A survey on coverage path planning for robotics," Robotics and Autonomous Systems, 第 61 卷, 第 12 期, 頁 1258–1276, 2013.',
     'https://www.sciencedirect.com/science/article/abs/pii/S092188901300167X'),
    ('H. Choset and P. Pignon, "Coverage path planning: The boustrophedon cellular decomposition," in Field and Service Robotics, Springer, 1998, 頁 203–209.',
     'https://www.ri.cmu.edu/pub_files/pub4/choset_howie_1997_3/choset_howie_1997_3.pdf'),
    ('G. Grisetti, C. Stachniss, and W. Burgard, "Improved techniques for grid mapping with Rao-Blackwellized particle filters," IEEE Transactions on Robotics, 第 23 卷, 第 1 期, 頁 34–46, 2007.',
     'https://www.ipb.uni-bonn.de/wp-content/papercite-data/pdf/grisetti07tro.pdf'),
    ('D. Fox, W. Burgard, and S. Thrun, "The dynamic window approach to collision avoidance," IEEE Robotics & Automation Magazine, 第 4 卷, 第 1 期, 頁 23–33, 1997.',
     'https://www.ri.cmu.edu/publications/the-dynamic-window-approach-to-collision-avoidance/'),
    ('M. Quigley et al., "ROS: an open-source Robot Operating System," in ICRA Workshop on Open Source Software, 2009.',
     'http://ai.stanford.edu/~mquigley/papers/icra2009-ros.pdf'),
    ('I. T. Jolliffe and J. Cadima, "Principal component analysis: a review and recent developments," Phil. Trans. R. Soc. A, 第 374 卷, 第 2065 期, 2016.',
     'https://doi.org/10.1098/rsta.2015.0202'),
    ('R. Bormann, F. Jordan, J. Hampp, and M. Hagele, "Indoor coverage path planning: Survey, implementation, analysis," in Proc. IEEE ICRA, 2018, 頁 1718–1725.',
     'https://ieeexplore.ieee.org/document/8460991'),
    ('ROBOTIS, "TurtleBot3 e-Manual." (TurtleBot3 官方文件)',
     'https://emanual.robotis.com/docs/en/platform/turtlebot3/overview/'),
    ('ROS Wiki, "move_base." (ROS 導航堆疊文件)',
     'http://wiki.ros.org/move_base'),
]


# ════════════════════════════════════════════════════════════
#  文件主體
# ════════════════════════════════════════════════════════════
def build():
    doc = base_doc()

    # ── 標題與作者 ──
    title(doc, "從零開始理解一套機器人完全覆蓋路徑規劃系統：")
    title(doc, "以 ROS、PCA 對齊牛耕式分解與即時網頁視覺化之 TurtleBot3 實作為例")
    center(doc, "（學習導向教學論文）", size=11, italic=True, after=8)
    center(doc, "MXHHulk", size=11, bold=True)
    center(doc, "TurtleBot3 完全覆蓋路徑規劃專案（rebuild 分支）", size=10, italic=True)
    center(doc, "hulk71777@gmail.com", size=10, italic=True, after=10)

    # ── 摘要 ──
    abstract(doc, "摘要——",
        "本文是一份「學習導向」的教學論文，目的並非僅報告成果，而是帶領一位「完全沒有"
        "機器人或 ROS 背景」的讀者，由淺入深地讀懂一整套可實際運行的完全覆蓋路徑規劃"
        "（Complete Coverage Path Planning, CCPP）系統。所謂完全覆蓋，是指讓移動機器人"
        "走遍空間中每一個可到達之處，最典型的例子就是掃地機器人。本系統建構於機器人作業"
        "系統（ROS Noetic）之上，以 TurtleBot3 Burger 為硬體平台，串接 gmapping 同時定位"
        "與建圖（SLAM）、牛耕式（來回往復）覆蓋規劃、move_base 導航與動態視窗法（DWA）"
        "避障，並提供一個以 Flask 為基礎的即時網頁監控介面。系統的核心技術亮點為：以主成分"
        "分析（PCA）自動將掃描方向對齊房間主軸，從而避免地圖傾斜所造成的階梯狀鋸齒路徑；"
        "並將演算法與 ROS 徹底解耦，形成可獨立測試、跨節點共用的純演算法模組。本文先以大篇幅"
        "鋪陳必要的預備知識（ROS 的節點與話題、SLAM、佔據網格、座標系、形態學、PCA），再"
        "逐章拆解系統架構、演算法、實作細節、併發安全、操作流程與除錯，並輔以程式碼片段、"
        "表格與示意圖，務求使初學者能完整建立全貌。")
    abstract(doc, "關鍵詞——",
        "完全覆蓋路徑規劃；牛耕式；主成分分析；ROS；SLAM；佔據網格；move_base；"
        "動態視窗法；多執行緒；網頁視覺化；TurtleBot3；教學")

    # ── 如何閱讀本文 ──
    h2(doc, "閱讀指引")
    body(doc,
        "本文章節依「由簡到難」排列。第 1 節說明這個專案到底在做什麼、解決什麼問題；第 2 節"
        "是給零基礎讀者的預備知識速成，若你已熟悉 ROS 可略過；第 3 節給出系統全貌與資料流；"
        "第 4 節是全文核心——覆蓋演算法（含 PCA 對齊）；第 5 節深入四支程式的實作；第 6 至 8 節"
        "處理座標轉換、併發安全與地圖圖層等較進階的工程細節；第 9 節教你如何實際啟動與操作；"
        "第 10 節整理常見問題與除錯；第 11 節提供延伸練習；第 12 節結論。文末附參考文獻、檔案"
        "地圖、參數總表與名詞對照表，可隨時翻查。")

    pagebreak(doc)

    # ════════════════════════════════════════════════
    #  1. 緒論
    # ════════════════════════════════════════════════
    h1(doc, "1　緒論：這個專案在解決什麼問題")

    h2(doc, "1.1　從掃地機器人說起")
    body(doc,
        "想像家裡的掃地機器人。它的任務和一般導航（從客廳走到廚房）有個根本差別：它不是要"
        "「從 A 點走到 B 點」，而是要「把整片地板都掃過一遍」，而且最好不重複、不漏掉。這種"
        "「走遍每一個角落」的問題，在機器人學裡稱為完全覆蓋路徑規劃（Complete Coverage Path "
        "Planning，簡稱 CCPP）[1]。除了掃地，地面巡檢、割草、噴漆、農田作業，本質上都是同一"
        "類問題。")
    body(doc,
        "本專案要做的，就是讓一台 TurtleBot3 Burger 小型機器人，在一個封閉空間（例如一個房間）"
        "裡自動走遍全場，並且全程可以在網頁上即時觀看：地圖長怎樣、機器人現在在哪、走過哪些"
        "地方、規劃的覆蓋路線是什麼，還能用按鈕一鍵「開始 / 停止」。")

    h2(doc, "1.2　覆蓋策略：什麼是「牛耕式」")
    body(doc,
        "本專案採用的覆蓋策略稱為牛耕式（Boustrophedon）[2]。這個字源自古希臘文「像牛犁田一樣」"
        "——牛拉著犁，一條直線耕到底，轉個彎，再平行耕回來，來回往復。把這個走法套到機器人身上，"
        "就是「掃一條直線到底，轉彎，平行掃回來」，像下圖這樣的 S 形：")
    code(doc,
        "  →  →  →  →  →  →  →\n"
        "                     ↓\n"
        "  ←  ←  ←  ←  ←  ←  ←\n"
        "  ↓\n"
        "  →  →  →  →  →  →  →")
    caption(doc, "圖 1　牛耕式（S 形來回）覆蓋示意。")

    h2(doc, "1.3　本專案的三項技術亮點（貢獻）")
    body(doc, "若用學術論文的語言來說，本專案有三個值得一提的設計重點：")
    numbered(doc,
        "PCA 自動對齊：SLAM 建出的地圖，牆面常與座標軸歪斜；若硬沿座標軸來回掃，路徑會變成"
        "破碎的階梯狀。本專案用主成分分析（PCA）找出房間自己的主軸方向，讓掃描線自動與牆面"
        "平行，大幅減少轉彎、提升覆蓋品質。")
    numbered(doc,
        "演算法與 ROS 解耦：把覆蓋演算法抽成一個完全不依賴 ROS 的純 Python 模組"
        "（coverage_planner.py），因此可以單獨測試，也能被多個節點共用。這是很好的工程習慣。")
    numbered(doc,
        "即時網頁視覺化：用 Flask + HTML5 Canvas 做出多圖層即時監控介面，把地圖、機器人、"
        "軌跡、覆蓋路徑疊在一起顯示，並提供一鍵任務控制。")

    # ════════════════════════════════════════════════
    #  2. 預備知識
    # ════════════════════════════════════════════════
    h1(doc, "2　預備知識速成（零基礎可讀）")
    body(doc,
        "要看懂本專案，需要先認識幾個機器人領域的觀念。本節盡量用生活化的比喻說明；若你已熟悉 "
        "ROS，可直接跳到第 3 節。下表先給一個總覽，後續小節再逐一展開。")
    table(doc,
        ["名詞", "一句話解釋"],
        [
            ["ROS (Noetic)", "機器人作業系統；本質是讓許多小程式（節點）彼此傳訊息的框架。"],
            ["節點 Node", "一支獨立執行的程式。本專案有 map_server、move_base、gmapping 等。"],
            ["話題 Topic", "節點之間傳訊息的「頻道」，發布者 publish、訂閱者 subscribe。"],
            ["訊息 Message", "在話題上流動的資料，型別事先定好（如 OccupancyGrid、Odometry）。"],
            ["/map", "一張地圖，型別為 OccupancyGrid（佔據網格）。"],
            ["/odom", "里程計，機器人自己估計的位置，型別 Odometry。"],
            ["/scan", "360° 光達掃描資料，型別 LaserScan。"],
            ["SLAM / gmapping", "邊走邊建地圖 + 同時定位；本專案用 gmapping 產生 /map。"],
            ["OccupancyGrid", "把空間切成格子，每格存 0=空地、100=障礙、-1=未知。"],
            ["move_base", "ROS 導航堆疊；給它目標點，它自己避障走過去。"],
            ["DWA", "move_base 內部的局部避障演算法（動態視窗法）。"],
            ["TF", "座標轉換系統，記錄各座標系（map、odom、機器人）之間的關係。"],
        ],
        widths=[1.6, 5.0])

    h2(doc, "2.1　ROS 的核心觀念：節點 + 話題 + 訊息")
    body(doc,
        "ROS 最關鍵的想法是：不要把整個機器人寫成一支大程式，而是拆成許多小程式（節點），"
        "它們彼此靠「傳訊息」溝通。可以用 YouTube 來比喻：話題像「頻道」，發布者是「上傳影片"
        "的人」，訂閱者是「按了訂閱的人」。發布者與訂閱者彼此不需要認識，只要約好頻道名稱與"
        "影片格式（訊息型別）就能溝通。這種「解耦」正是 ROS 能輕鬆組合不同模組的關鍵。")
    body(doc,
        "每個節點啟動時都會先向系統「報到」。例如 map_server.py 的第一行就是註冊自己的名字：")
    code(doc, 'rospy.init_node(\'map_server\')   # 向 ROS 註冊「我叫 map_server」')

    h3(doc, "2.1.1　發布（Publish）")
    body(doc,
        "「發布」就是把一則訊息丟到某個話題上，誰想看誰自己訂閱。以 boustrophedon.py 為例，"
        "它算出覆蓋路徑後，發布到 /coverage_path 供 RViz 顯示。先建立發布者，再送出訊息：")
    code(doc,
        "# 第 1 步：開一個發布者\n"
        "_pub = rospy.Publisher('/coverage_path', Path, queue_size=1, latch=True)\n"
        "#                       └ 話題名稱       └ 訊息型別 └ 佇列   └ 鎖存\n"
        "\n"
        "# 第 2 步：把訊息送出去\n"
        "path = Path()\n"
        "path.header.frame_id = 'map'\n"
        "path.poses = [_make_pose(x, y, 'map') for x, y in pts]\n"
        "_pub.publish(path)        # ← 真正發布（按下上傳鍵）")
    body(doc,
        "其中 latch=True（鎖存）很實用：最後一則訊息會被「留住」，之後才上線的訂閱者（例如晚一點"
        "才打開的 RViz）也能立刻收到上一次的路徑，不必空等下一次重算。")

    h3(doc, "2.1.2　訂閱（Subscribe）與回呼函式")
    body(doc,
        "「訂閱」就是跟 ROS 說：「這個話題只要一有新訊息，就幫我呼叫某個函式處理。」這個被呼叫"
        "的函式稱為回呼函式（callback）。map_server.py 啟動時一口氣訂閱兩個話題：")
    code(doc,
        "rospy.Subscriber('/map',  OccupancyGrid, map_callback,  queue_size=1)\n"
        "rospy.Subscriber('/odom', Odometry,      odom_callback, queue_size=10)\n"
        "#                 └ 話題    └ 訊息型別      └ 收到就呼叫誰  └ 佇列長度")
    body(doc,
        "/odom 更新很快，佇列留 10 則避免漏接；/map 又大又慢，只留最新 1 則即可。訂閱者完全不需要"
        "知道是誰在發布 /map（其實是 gmapping），哪天換成別的建圖工具，只要對方一樣發布 /map，"
        "這支訂閱程式一行都不用改——這就是 pub/sub 解耦的威力。最後要讓程式「掛著」持續等訊息：")
    code(doc, "rospy.spin()   # 卡住主執行緒，把 CPU 交給 ROS 去觸發各個回呼")

    h3(doc, "2.1.3　另一種溝通：Action（動作）")
    body(doc,
        "並非所有溝通都走 pub/sub。本專案把目標點交給 move_base 時用的是 Action 機制，它適合"
        "「有開始、要等結果、可中途取消」的長時間任務（例如「走到這個點」）。pub/sub 適合持續"
        "廣播的資料流，Action 適合一次性的任務委派。第 5.2 與第 7 節會再深入。")

    h2(doc, "2.2　SLAM 與 gmapping：邊走邊畫地圖")
    body(doc,
        "機器人剛開機時對環境一無所知。SLAM（Simultaneous Localization and Mapping，同時定位"
        "與建圖）讓它一邊用光達掃描周遭、一邊推算自己的位置，逐步把地圖拼出來。本專案使用 "
        "gmapping [3]，它吃 /scan 光達資料，輸出兩樣東西：(1) /map 佔據網格地圖；(2) map→odom 的"
        "TF 座標轉換（讓 move_base 知道機器人在地圖上的位置）。")

    h2(doc, "2.3　佔據網格（OccupancyGrid）")
    body(doc,
        "佔據網格是地圖的資料結構。想像一張方格紙鋪在房間地板上，每一格記錄那塊地是什麼狀態。"
        "整張地圖就是一個二維陣列，每格存一個整數：")
    table(doc,
        ["數值", "意義", "灰階顯示"],
        [
            ["0", "自由（空地，可走）", "白 (255)"],
            ["100", "障礙（牆、家具）", "黑 (0)"],
            ["-1", "未知（還沒探索過）", "灰 (128)"],
        ],
        widths=[1.2, 3.2, 2.2])
    body(doc,
        "除了格子數值，訊息還附帶「中繼資料」：resolution（解析度，每格幾公尺，本專案約 0.05 m）、"
        "origin（地圖原點的世界座標）、width / height（地圖寬高，單位是格）。這些之後做座標轉換"
        "時都會用到。")

    h2(doc, "2.4　三種座標系")
    body(doc,
        "本專案會在三種座標之間轉換，先建立概念，第 6 節會給出實際公式：")
    table(doc,
        ["座標系", "說明", "單位"],
        [
            ["世界座標 (map)", "gmapping 的全域座標，路徑規劃的基準", "公尺"],
            ["網格座標 (grid)", "OccupancyGrid 的行 (row)、列 (col) 索引", "格"],
            ["圖片像素 (pixel)", "裁切後顯示在網頁 Canvas 上的像素位置", "像素"],
        ],
        widths=[1.8, 3.6, 1.2])

    h2(doc, "2.5　move_base 與 DWA（不是同一層的東西）")
    body(doc,
        "這兩個名詞常被混淆，但其實是「外包總管 vs 公司裡某個部門」的關係。move_base 是整套導航的"
        "總管框架 [9]：你丟一個目標點給它，它負責「想辦法安全開過去」這整件事。DWA（動態視窗法，"
        "Dynamic Window Approach）[4] 則是 move_base 內部負責「即時避障、決定下一步馬達怎麼轉」"
        "的那個小元件（局部規劃器）。換句話說，DWA 是 move_base 裡的一個零件，不是平行的兩個東西。")
    body(doc, "move_base 內部其實同時運作好幾塊：")
    table(doc,
        ["內部元件", "工作"],
        [
            ["Global planner（全域規劃器）", "看整張地圖，規劃從現在到目標的大致路線。"],
            ["Local planner（局部規劃器）", "就是 DWA。盯著眼前幾秒，即時決定速度/轉向、閃避障礙。"],
            ["Global costmap", "整張地圖的障礙＋膨脹層，給全域規劃器用。"],
            ["Local costmap", "機器人周圍一小塊的即時障礙圖（吃光達），給 DWA 用。"],
            ["Recovery behaviors", "卡住時的補救動作：原地旋轉、清除 costmap 等。"],
        ],
        widths=[2.4, 4.2])
    body(doc,
        "DWA 的運作可濃縮為五步：(1) 在「當前速度 + 馬達加速度上限」可達範圍內（這就是「動態視窗」）"
        "取樣一堆候選速度 (v, ω)；(2) 模擬每個組合接下來一小段會走到哪；(3) 用評分函式打分（離障礙"
        "夠遠嗎？有朝全域路線前進嗎？夠快嗎？）；(4) 選最高分輸出成 /cmd_vel 速度指令給馬達；"
        "(5) 每幾十毫秒重複一次，因此能即時對突然出現的障礙反應。記憶口訣：問「要不要去、去哪、"
        "到了沒」是 move_base 的層級；問「這一瞬間馬達該轉多快」是 DWA 的層級。")

    h2(doc, "2.6　數學預備一：主成分分析（PCA）")
    body(doc,
        "PCA（Principal Component Analysis）[6] 是一種找出「一堆點主要朝哪個方向分布」的方法。"
        "想像把一把米粒灑在桌上排成橢圓形，PCA 能告訴你這個橢圓的長軸與短軸方向。做法是：把所有點"
        "的座標算出共變異數矩陣，再求它的特徵值與特徵向量；最大特徵值對應的特徵向量就是「分布最開"
        "的方向」（長軸），最小的就是短軸。本專案用它找出房間自由空間的主軸，讓掃描線沿長軸走。"
        "第 4 節會詳述。")

    h2(doc, "2.7　數學預備二：形態學（侵蝕與膨脹）")
    body(doc,
        "形態學（morphology）是影像處理裡操作形狀的工具。對黑白圖中的「障礙」區域：膨脹（dilation）"
        "會把障礙往外長胖一圈；侵蝕（erosion）會把障礙往內縮小一圈。本專案用膨脹來製造安全邊距"
        "——把牆「加胖」機器人半徑那麼多，規劃時只要避開加胖後的牆，機器人本體（有體積）就不會"
        "撞牆。第 5 與第 8 節會說明三種地圖圖層如何由此而來。")

    pagebreak(doc)

    # ════════════════════════════════════════════════
    #  3. 系統架構
    # ════════════════════════════════════════════════
    h1(doc, "3　系統架構：資料如何流動")
    body(doc,
        "有了預備知識，現在看全貌。整套系統由一組互相協作的 ROS 節點構成，資料流如下圖：")
    code(doc,
        "┌───────────────────────────────────────────────┐\n"
        "│              實體 TurtleBot3 Burger             │\n"
        "│           （360° 光達 LDS-01 + 馬達）           │\n"
        "└──────────────┬──────────────────┬──────────────┘\n"
        "        /scan   │                  │  /odom\n"
        "                ▼                  │\n"
        "        ┌───────────────┐         │\n"
        "        │ gmapping SLAM │         │\n"
        "        │ 邊走邊建圖+定位│         │\n"
        "        └───────┬───────┘         │\n"
        "    /map(網格)   │ + TF(map→odom)  │\n"
        "  ┌─────────────┼──────────────────┼─────────────┐\n"
        "  │             ▼                  ▼              │\n"
        "  │   ┌───────────────┐   ┌─────────────────┐    │\n"
        "  │   │boustrophedon.py│   │  map_server.py  │    │\n"
        "  │   │ (模組B,選用)   │   │   (模組A,核心)  │    │\n"
        "  │   │算路徑→         │   │地圖圖層+網頁+   │    │\n"
        "  │   │ /coverage_path │   │  派送任務       │    │\n"
        "  │   └───────┬───────┘   └────────┬────────┘    │\n"
        "  │           │ 共用            目標點│            │\n"
        "  │   ┌───────▼─────────┐  ┌────────▼────────┐   │\n"
        "  │   │coverage_planner │  │    move_base    │   │\n"
        "  │   │(純演算法,無 ROS)│  │ 導航 + DWA 避障 │   │\n"
        "  │   └─────────────────┘  └────────┬────────┘   │\n"
        "  └──────────────────────────────────┼───────────┘\n"
        "                              /cmd_vel │\n"
        "                                       ▼\n"
        "                                    馬達轉動")
    caption(doc, "圖 2　系統資料流總覽。實線箭頭為 ROS 話題/動作；coverage_planner 被兩個模組共用。")
    body(doc,
        "一句話流程：光達掃描 → gmapping 建出 /map → 規劃模組把 /map 變成一串覆蓋路點 → "
        "map_server 把路點一個個丟給 move_base → move_base 開著機器人避障走過去 → 網頁即時"
        "顯示全部過程。下表整理各組成的角色：")
    table(doc,
        ["組成", "角色"],
        [
            ["TurtleBot3（硬體）", "提供 /scan 光達與 /odom 里程計，接收 /cmd_vel 驅動馬達。"],
            ["gmapping", "SLAM；輸出 /map 與 map→odom 的 TF。"],
            ["coverage_planner.py", "純演算法核心（PCA + 牛耕），無 ROS 依賴，被兩個模組共用。"],
            ["boustrophedon.py（模組 B）", "規劃節點；發布 /coverage_path 供 RViz 顯示（預設停用）。"],
            ["map_server.py（模組 A）", "核心樞紐；地圖圖層、網頁伺服器、軌跡記錄、派送任務。"],
            ["move_base", "導航；接收目標點，內部用 DWA 避障，輸出 /cmd_vel。"],
            ["web/index.html", "前端；Canvas 多圖層即時監控與一鍵控制。"],
        ],
        widths=[2.3, 4.3])

    h2(doc, "3.1　哪些節點是「我寫的」，哪些來自外部套件")
    body(doc,
        "初學者常見的疑惑是：上圖那麼多方塊，到底哪些是本專案的程式碼？答案是——大部分都不是。"
        "本專案真正自己撰寫的只有「模組 A／B」兩支腳本，其餘像 gmapping、move_base、硬體驅動，"
        "全都是成熟的開源套件，本專案只是把它們「組裝」起來並下達指令。這正是 ROS 的價值：站在"
        "巨人的肩膀上，不必重造輪子。下表清楚標示每個節點的歸屬與來源：")
    table(doc,
        ["節點 / 啟動項", "由誰啟動", "來源套件", "歸屬"],
        [
            ["map_server", "start.launch（本專案）", "turtlebot3_ccpp", "★ 自製"],
            ["boustrophedon_planner（預設停用）", "start.launch（已註解）", "turtlebot3_ccpp", "★ 自製"],
            ["coverage_planner（非節點，是模組）", "被上兩者 import", "turtlebot3_ccpp", "★ 自製"],
            ["硬體驅動（發布 /scan、/odom，接收 /cmd_vel）", "turtlebot3_robot.launch", "turtlebot3_bringup", "外部（ROBOTIS 官方）"],
            ["slam_gmapping（建 /map 與 map→odom TF）", "turtlebot3_slam.launch", "gmapping", "外部（OpenSLAM / ROS）"],
            ["move_base（導航總管）", "start.launch（本專案啟動，參數沿用官方）", "move_base", "外部（ROS Navigation Stack）"],
            ["DWA 局部規劃器、costmap", "move_base 內部載入", "dwa_local_planner、costmap_2d", "外部（Navigation Stack）"],
        ],
        widths=[2.3, 1.9, 1.6, 1.4])
    body(doc,
        "也就是說：本專案自己寫的節點只有 map_server（核心）與 boustrophedon_planner（選用，預設"
        "停用）；coverage_planner 則是被它們共用的純演算法「模組」，本身不是一個獨立節點。其餘節點"
        "都來自三個外部套件家族：(1) ROBOTIS 官方的 TurtleBot3 套件（turtlebot3_bringup 提供硬體"
        "驅動、turtlebot3_slam 包裝 gmapping、turtlebot3_navigation 提供導航參數檔）；(2) ROS 官方"
        "的導航堆疊 Navigation Stack（move_base、dwa_local_planner、costmap_2d）；(3) SLAM 社群的"
        "gmapping。")

    h3(doc, "3.1.1　這些外部節點是怎麼被「找到」並啟動的")
    body(doc,
        "關鍵在 start.launch 裡的 $(find 套件名) 語法。它請 ROS 在系統已安裝的套件中，找出該套件的"
        "實際安裝路徑，再去 include 它的 launch 檔或載入它的參數。例如本專案的 start.launch 就是這樣"
        "把三個外部套件「接」進來：")
    code(doc,
        "<include file=\"$(find turtlebot3_bringup)/launch/turtlebot3_robot.launch\" />\n"
        "<include file=\"$(find turtlebot3_slam)/launch/turtlebot3_slam.launch\">\n"
        "    <arg name=\"slam_methods\" value=\"gmapping\" />\n"
        "</include>\n"
        "<node pkg=\"move_base\" type=\"move_base\" name=\"move_base\" ...>\n"
        "    <rosparam file=\"$(find turtlebot3_navigation)/param/...yaml\" command=\"load\" />\n"
        "</node>")
    body(doc,
        "因此本專案的 start.launch 本質上是一份「組裝清單」：先請出官方硬體驅動，再請出 gmapping "
        "建圖，再啟動 move_base 並餵入官方導航參數，最後才掛上自製的 map_server。換套件、換型號"
        "（model:=waffle）都只是改這份清單，不必動到核心程式。")

    h3(doc, "3.1.2　這些外部套件實際裝在哪、如何查證")
    body(doc,
        "這些套件通常透過下列兩種方式之一安裝在機器上：(1) 以 apt 安裝官方二進位套件（檔案位於 "
        "/opt/ros/noetic/share/ 之下），套件名多為 ros-noetic-xxx；(2) 從 ROBOTIS／ros-planning 的 "
        "GitHub 原始碼複製到 catkin 工作區（~/catkin_ws/src）後自行編譯。你可以用以下指令親自查證"
        "每個節點究竟來自哪個套件、安裝在哪個路徑：")
    code(doc,
        "rospack find turtlebot3_bringup     # 印出該套件的安裝路徑\n"
        "rospack find gmapping               # 同理，找 gmapping\n"
        "rospack find move_base              # 找 move_base\n"
        "roscd turtlebot3_navigation/param   # 直接跳到官方導航參數檔目錄\n"
        "rosnode info /move_base             # 看某個執行中節點的話題與資訊\n"
        "apt list --installed | grep -E 'turtlebot3|gmapping|navigation'  # 看 apt 裝了哪些\n"
        "dpkg -S $(rospack find gmapping)    # 反查某路徑屬於哪個 apt 套件")
    body(doc,
        "簡言之：凡是出現在 scripts/ 目錄下的（map_server.py、boustrophedon.py、coverage_planner.py）"
        "才是本專案的程式碼；凡是用 $(find ...) 引用的，都是別人寫好、由社群維護的外部套件。理解"
        "這條界線，就能分清楚「哪裡該改自己的程式、哪裡該查官方文件」。")

    # ════════════════════════════════════════════════
    #  4. 核心演算法
    # ════════════════════════════════════════════════
    h1(doc, "4　核心演算法：牛耕式 + PCA 對齊")
    body(doc,
        "本節是全文核心，所有規劃邏輯都集中在不依賴 ROS 的 coverage_planner.py。它只用 NumPy 與 "
        "SciPy，輸入輸出都是普通陣列與座標，提供兩個函式：apply_safety_margin（安全邊距膨脹）與 "
        "boustrophedon（牛耕路點生成）。兩個關鍵參數如下表（取自當前程式碼）：")
    table(doc,
        ["參數", "值", "意義"],
        [
            ["SPACING", "0.18 m", "掃描線間距；覆蓋寬度約 0.20 m，留約 10% 重疊不漏掃。"],
            ["MARGIN", "0.10 m", "安全邊距，約等於機器人半徑（Burger 直徑約 0.20 m）。"],
        ],
        widths=[1.4, 1.2, 4.0])

    h2(doc, "4.1　第一步：自由空間擷取與安全邊距")
    body(doc,
        "規劃前要先決定「哪裡可以走」。直接用 data==0（空地）還不夠，因為機器人有體積，貼著牆走"
        "會撞牆。所以先對障礙物膨脹一個機器人半徑，再把這圈安全區從可走區扣掉。apply_safety_margin "
        "正是做這件事：把障礙（值 100）用一個邊長 2r+1 的方形核做膨脹，其中 r = round(MARGIN/解析度)。")
    code(doc,
        "def apply_safety_margin(data, margin, resolution):\n"
        "    obs = data == 100                       # 找出障礙格\n"
        "    r = max(1, round(margin / resolution))  # 邊距換算成格數\n"
        "    kern = np.ones((2*r+1, 2*r+1), dtype=bool)\n"
        "    return binary_dilation(obs, structure=kern)   # 障礙往外長胖 r 格")
    body(doc, "有了膨脹後的障礙，可走集合就定義為「是空地、且不在膨脹障礙內」：")
    code(doc,
        "safe_obs = apply_safety_margin(data, MARGIN, resolution)\n"
        "free     = (data == 0) & ~safe_obs     # 可走 = 空地 且 非安全邊距區")

    h2(doc, "4.2　第二步：為什麼需要 PCA")
    body(doc,
        "牛耕式最簡單的做法是「沿著 X 軸來回掃」。但問題是：SLAM 建出來的地圖，房間牆壁往往不是"
        "正對 X/Y 軸的，而是歪斜一個角度（因為機器人開機那一刻的朝向決定了地圖座標方向）。如果"
        "牆是斜的、掃描線是正的，覆蓋路徑就會變成鋸齒/階梯狀，效率差、轉彎多：")
    code(doc,
        "  牆是斜的，但掃描線是水平的 → 邊緣全是碎階梯\n"
        "     ╱──────────╱\n"
        "    ╱  → → → →  ╱\n"
        "   ╱  → → → →  ╱\n"
        "  ╱──────────╱")
    caption(doc, "圖 3　固定軸向掃描遇上傾斜房間，產生階梯狀鋸齒。")
    body(doc,
        "解法是用 PCA 找出房間自己的方向。具體做法（對應 coverage_planner.py 第 49–62 行）："
        "(1) 取出所有可走格的世界座標，得到一堆點；(2) 計算這堆點的共變異數矩陣，求特徵值與特徵"
        "向量（np.linalg.eigh）；(3) 最大特徵值對應的向量＝空間長軸（沿它掃，掃描線最長、轉彎"
        "最少）；(4) 最小特徵值對應的向量＝短軸（每掃完一條線，沿短軸前進一個 spacing）。這樣"
        "不管地圖怎麼歪，掃描線都會自動和牆壁平行。")
    code(doc,
        "center           = pts.mean(axis=0)\n"
        "diffs            = pts - center\n"
        "eigvals, eigvecs = np.linalg.eigh(np.cov(diffs.T))\n"
        "axis_a = eigvecs[:, np.argmax(eigvals)]   # sweep 方向（長軸）\n"
        "axis_b = eigvecs[:, np.argmin(eigvals)]   # step  方向（短軸）")

    h2(doc, "4.3　第三步：牛耕路點怎麼生成")
    body(doc,
        "有了長短軸，就把每個可走點投影到長軸（proj_a）與短軸（proj_b）。然後沿短軸從一端走到"
        "另一端，每隔 spacing 取一條「掃描線」；每條線上找出最遠的兩個端點連成線段；再用一個布林"
        "旗標 l2r 控制方向交替——這條左到右，下條就右到左，正是牛耕來回的精髓。")
    code(doc,
        "b = b_min + spacing / 2     # 第一條線內縮半個間距，兩端對稱覆蓋\n"
        "l2r = True\n"
        "while b <= b_max + spacing / 2:\n"
        "    mask = np.abs(proj_b - b) < spacing / 2     # 落在這條線附近的點\n"
        "    if mask.any():\n"
        "        a_vals  = proj_a[mask]\n"
        "        p_start = center + a_vals.min()*axis_a + b*axis_b\n"
        "        p_end   = center + a_vals.max()*axis_a + b*axis_b\n"
        "        if l2r: waypoints += [p_start, p_end]\n"
        "        else:   waypoints += [p_end,   p_start]   # 反向\n"
        "        l2r = not l2r          # 下一條反過來\n"
        "    b += spacing")
    body(doc,
        "注意第一條掃描線從邊界內縮 spacing/2（b = b_min + spacing/2），讓兩端對稱、覆蓋更均勻。"
        "最終輸出是一串依序走訪的世界座標路點 [(x, y), ...]。")

    h2(doc, "4.4　第四步：執行與失敗容錯")
    body(doc,
        "規劃出路點後，map_server 逐一把它們送給 move_base，並讓每個目標朝向下一個路點（用 "
        "atan2 算偏航角再轉成四元數）。move_base 內部用 DWA 規劃全域路線並即時避障。關鍵的"
        "工程細節是失敗容錯：若某個路點到不了（move_base 回傳非 SUCCEEDED 狀態），系統不會卡死，"
        "而是記錄警告後跳過、繼續下一點，確保單一無法到達的目標不會中止整個任務。")
    code(doc,
        "if client.get_state() != GoalStatus.SUCCEEDED:\n"
        "    rospy.logwarn(f'[coverage] 路點 {i+1}/{n} 跳過（狀態 {client.get_state()}）')\n"
        "    cov_status['msg'] = f'路點 {i+1}/{n} 跳過（無法到達）'\n"
        "    # 不 return，繼續 for 迴圈下一個路點")

    pagebreak(doc)

    # ════════════════════════════════════════════════
    #  5. 實作細節
    # ════════════════════════════════════════════════
    h1(doc, "5　實作細節：四支程式逐一拆解")
    body(doc,
        "本專案刻意把演算法與 ROS 接線拆開：三支 Python 腳本 + 一個網頁 + 一個 launch。下面逐一說明。")

    h2(doc, "5.1　coverage_planner.py —— 純演算法核心（模組共用）")
    bullet(doc, "只用 numpy 與 scipy，輸入輸出都是普通陣列與座標，完全不 import 任何 ROS。")
    bullet(doc, "因為不依賴 ROS，所以可以單獨測試，也能被 boustrophedon.py 與 map_server.py 共用。")
    bullet(doc, "提供 apply_safety_margin() 與 boustrophedon() 兩個函式（已於第 4 節詳述）。")
    bullet(doc, "定義兩個關鍵常數 SPACING=0.18、MARGIN=0.10，成為全系統的「單一真相來源」。")

    h2(doc, "5.2　map_server.py —— 地圖伺服器 + 任務執行（核心樞紐）")
    body(doc, "這支最複雜，身兼四職：")
    numbered(doc,
        "訂閱 /map（map_callback）：把佔據網格轉成三種 PNG 圖層（原始/侵蝕/膨脹），並裁切掉"
        "周圍空白（CROP_PAD=10 格邊距）。")
    numbered(doc,
        "訂閱 /odom（odom_callback）：記錄機器人位置與軌跡，每移動 0.1 m 記一點，最多存 10000 點。")
    numbered(doc,
        "Flask 網頁伺服器：開在 8080 埠，提供 /map.png、/robot_state、/coverage/start|stop|status 等 API。")
    numbered(doc,
        "任務執行緒（run_coverage）：把覆蓋路點一個個送給 move_base，監控到達狀態，可隨時被停止。")
    body(doc,
        "它透過 ROS 的 Action 介面連上 move_base。注意它只負責「派工」，完全碰不到也不需要知道"
        "DWA 的細節：")
    code(doc,
        "client = actionlib.SimpleActionClient('move_base', MoveBaseAction)\n"
        "...\n"
        "goal = MoveBaseGoal()\n"
        "goal.target_pose.header.frame_id = fid\n"
        "goal.target_pose.pose.position.x = x\n"
        "goal.target_pose.pose.position.y = y\n"
        "goal.target_pose.pose.orientation.z, ...w = _yaw_to_quat(yaw)\n"
        "client.send_goal(goal)        # 「這個點，你想辦法開過去」")
    body(doc,
        "這裡也展現了「網頁所見即規劃所避」的單一真相設計：網頁顯示的膨脹層，直接呼叫"
        "與規劃同一個 apply_safety_margin 函式產生，保證使用者在網頁上看到的安全邊距，"
        "就是機器人實際避開的區域。")

    h2(doc, "5.3　boustrophedon.py —— 規劃節點（模組 B，預設停用）")
    bullet(doc, "訂閱 /map，每隔 REPLAN_INTERVAL=5.0 秒重新規劃一次（節流，避免地圖頻繁更新時瘋狂重算）。")
    bullet(doc, "呼叫 coverage_planner 算出路點，包成 nav_msgs/Path，發布到 /coverage_path。")
    bullet(doc, "這個節點的路徑只給 RViz 顯示用，不驅動機器人；在 start.launch 中預設被註解停用，"
                "因為網頁端的覆蓋路徑已由 map_server 自行計算顯示，兩者功能重複。")

    h2(doc, "5.4　web/index.html —— 前端網頁")
    bullet(doc, "用 HTML5 Canvas 多圖層疊加：原始地圖打底，侵蝕/膨脹層著色半透明疊上，再畫覆蓋路徑、"
                "行走軌跡、機器人位置。")
    bullet(doc, "每秒輪詢一次後端（setInterval(refresh, 1000)），重新抓圖與狀態。")
    bullet(doc, "有趣的技巧 colorize()：把灰階地圖即時轉成彩色障礙圖層（空地透明、障礙著色），"
                "讓多層地圖能疊在一起看。")
    bullet(doc, "覆蓋路徑分「已走（實線）/ 待走（虛線）」；軌跡用漸亮淡出（近期亮、久遠淡）；"
                "機器人畫成帶光暈的圓點，半徑依解析度換算（約機器人實際大小）。")
    bullet(doc, "圖層可用按鈕一鍵開關，開始/停止按鈕對應 /coverage/start|stop。")

    h2(doc, "5.5　start.launch —— 一鍵啟動")
    body(doc, "把所有東西串起來，依序啟動：")
    numbered(doc, "設定環境變數 TURTLEBOT3_MODEL（預設 burger，可用 model:=waffle 切換）。")
    numbered(doc, "硬體驅動 turtlebot3_robot.launch，與實體機器人建立連線。")
    numbered(doc, "gmapping SLAM（turtlebot3_slam.launch，open_rviz=false）：發布 /map 與 map→odom TF。")
    numbered(doc, "move_base 導航：指定 DWA 局部規劃器，沿用 turtlebot3_navigation 的標準參數檔"
                  "（costmap、move_base、DWA 參數）。")
    numbered(doc, "map_server.py：開 8080 埠提供網頁監控與任務派送。（boustrophedon.py 預設註解停用。）")

    h2(doc, "5.6　測試小工具")
    body(doc,
        "test/ 目錄下有量測光達的小工具，例如 test_lidar_freq.py 訂閱 /scan、每 2 秒印出平均頻率"
        "（Hz），用來確認感測器是否正常運作；另有量測光達距離範圍的腳本。它們是獨立節點，與主系統"
        "解耦，方便除錯。")

    # ════════════════════════════════════════════════
    #  6. 座標轉換
    # ════════════════════════════════════════════════
    h1(doc, "6　座標系與轉換（進階細節）")
    body(doc,
        "網頁要把世界座標（公尺）畫到裁切後的圖片上（像素），需要一個轉換。核心公式在 "
        "map_server.py 的 world_to_px：")
    code(doc,
        "px = (wx - origin_x) / resolution - c0\n"
        "py = crop_h - ((wy - origin_y) / resolution - r0)")
    body(doc, "兩個重點：")
    numbered(doc, "（世界座標 − 原點）÷ 解析度 = 網格座標；再扣掉裁切左/上邊界 (c0, r0) 得到裁切後座標。")
    numbered(doc,
        "Y 軸要翻轉：影像座標原點在左上、Y 向下；世界座標 Y 向上。所以 py 用 crop_h 去減。"
        "（產生 PNG 時的 np.flipud 也是為了同一件事。）")

    # ════════════════════════════════════════════════
    #  7. 併發
    # ════════════════════════════════════════════════
    h1(doc, "7　併發與執行緒安全（進階細節）")
    body(doc, "map_server.py 同時跑著好幾條執行緒，必須小心資料競爭：")
    bullet(doc, "ROS 回呼執行緒：map_callback、odom_callback 不斷更新共享資料。")
    bullet(doc, "Flask 執行緒：處理網頁請求時要讀共享資料。")
    bullet(doc, "任務執行緒：run_coverage 在背景把路點送給 move_base。")
    body(doc, "解法是三把鎖，各自保護一塊共享狀態：")
    code(doc,
        "map_lock   = threading.Lock()   # 保護地圖資料\n"
        "robot_lock = threading.Lock()   # 保護機器人位置/軌跡\n"
        "cov_lock   = threading.Lock()   # 保護覆蓋任務狀態")
    body(doc,
        "每次讀寫共享資料都用 with xxx_lock: 包起來，確保 Flask 讀取時不會讀到寫到一半的資料。"
        "任務的停止機制也很值得學：run_coverage 不是被「強制 kill」，而是在迴圈裡定期檢查 "
        "cov_status['state']，發現變成非 running 就主動 cancel_all_goals() 並退出——這叫"
        "合作式取消（cooperative cancellation），安全又乾淨。")
    code(doc,
        "while not rospy.is_shutdown():\n"
        "    with cov_lock:\n"
        "        if cov_status['state'] != 'running':\n"
        "            client.cancel_all_goals()\n"
        "            return                       # 優雅退出\n"
        "    if client.wait_for_result(rospy.Duration(0.5)):\n"
        "        break                            # 每 0.5 秒檢查一次")

    # ════════════════════════════════════════════════
    #  8. 地圖圖層
    # ════════════════════════════════════════════════
    h1(doc, "8　地圖三種圖層：原始 / 侵蝕 / 膨脹")
    body(doc,
        "map_callback 會產生三種地圖，這是形態學影像處理的應用。灰階對應：白(255)=空地、"
        "黑(0)=障礙、灰(128)=未知。")
    table(doc,
        ["圖層", "端點", "做法", "用途"],
        [
            ["原始", "/map.png", "佔據網格直接轉灰階", "真實地圖，無形態學處理"],
            ["侵蝕", "/map_eroded.png", "binary_erosion 縮小障礙", "視覺參考（障礙縮小）"],
            ["膨脹", "/map_dilated.png", "apply_safety_margin 放大障礙", "規劃實際避開的安全邊距區"],
        ],
        widths=[0.9, 1.7, 2.4, 1.6])
    body(doc,
        "關鍵：膨脹層不是隨意放大，而是直接呼叫與路徑規劃同一個 apply_safety_margin（核半徑 = "
        "round(MARGIN/解析度)）。因此「網頁看到的膨脹」＝「機器人實際避開的區域」，是單一真相來源。"
        "侵蝕層則用相同尺寸的核做侵蝕，純為視覺參考，路徑規劃無對應物。")
    body(doc,
        "另外，地圖周圍常有一大片未知區(-1)。_crop 會找出已知格子的最小外接矩形，加上 CROP_PAD=10 "
        "格邊距後裁切，以減少網頁傳輸量、去掉多餘空白。")

    # ════════════════════════════════════════════════
    #  9. 操作
    # ════════════════════════════════════════════════
    h1(doc, "9　如何啟動與操作")
    h2(doc, "9.1　前置安裝")
    code(doc, "pip3 install flask Pillow      # numpy 已隨 ROS 安裝")
    h2(doc, "9.2　啟動")
    code(doc,
        "# 預設 Burger 型號\n"
        "roslaunch turtlebot3_ccpp start.launch\n"
        "\n"
        "# 換成 Waffle 型號\n"
        "roslaunch turtlebot3_ccpp start.launch model:=waffle\n"
        "\n"
        "# 換埠號\n"
        "roslaunch turtlebot3_ccpp start.launch port:=9000")
    h2(doc, "9.3　操作流程")
    numbered(doc, "啟動後，先遙控機器人在房間裡走一圈，讓 gmapping 把地圖建得夠完整。")
    numbered(doc, "瀏覽器打開 http://<機器人IP>:8080。")
    numbered(doc, "觀察地圖、機器人位置、規劃出來的覆蓋路徑（紫色）。")
    numbered(doc, "點「▶ 開始覆蓋」，機器人就會依路點自動走遍全場。")
    numbered(doc, "隨時可點「■ 停止」。")
    body(doc, "也可單獨查看某張地圖：/map.png（原始）、/map_eroded.png（侵蝕）、/map_dilated.png（膨脹）。")

    # ════════════════════════════════════════════════
    #  10. 除錯
    # ════════════════════════════════════════════════
    h1(doc, "10　常見問題與除錯")
    table(doc,
        ["症狀", "可能原因", "排查方向"],
        [
            ["網頁顯示「等待地圖...」", "/map 還沒發布", "確認 gmapping 啟動；rostopic echo /map 有資料"],
            ["「無可走路徑，地圖可能不完整」", "地圖太空或全是障礙", "先遙控機器人多建一點圖"],
            ["「move_base 未啟動」", "move_base 沒起來", "看 launch 輸出；rosnode list 確認"],
            ["覆蓋路徑呈鋸齒狀", "PCA 沒對齊（理論上罕見）", "檢查可走格數量是否太少"],
            ["機器人撞牆", "安全邊距太小", "調大 coverage_planner.py 的 MARGIN"],
            ["漏掃", "掃描線太疏", "調小 SPACING（但會變慢）"],
            ["路點一直跳過", "邊距太大堵死路，或定位漂移", "調 MARGIN；檢查 TF"],
            ["光達沒資料", "感測器/連線問題", "跑 test/test_lidar_freq.py 看頻率"],
        ],
        widths=[2.0, 2.0, 2.6])
    body(doc, "好用的除錯指令：")
    code(doc,
        "rostopic list                 # 看有哪些話題\n"
        "rostopic echo /map --noarr    # 看地圖 metadata（不印整個陣列）\n"
        "rostopic hz /scan             # 看光達頻率\n"
        "rosnode list                  # 看有哪些節點在跑\n"
        "rqt_graph                     # 視覺化節點與話題關係")

    # ════════════════════════════════════════════════
    #  11. 延伸練習
    # ════════════════════════════════════════════════
    h1(doc, "11　可延伸的練習")
    numbered(doc, "加一個「完全不裁切」的 /map_full.png 端點，比較裁切前後差異。")
    numbered(doc, "計算覆蓋率指標：走完後統計「實際走過的格 / 應覆蓋的格」，並顯示到網頁。")
    numbered(doc, "換覆蓋策略：試試螺旋式（spiral）或基於 frontier 的自主探索，省去人工建圖。")
    numbered(doc, "優化路點順序：目前是固定 S 形，可研究如何減少總路程（接近 TSP 問題）。")
    numbered(doc, "多區域分解：對非凸（L 形、多房間）環境，先切成子區域再各自牛耕。")

    # ════════════════════════════════════════════════
    #  12. 結論
    # ════════════════════════════════════════════════
    h1(doc, "12　結論")
    body(doc,
        "本文以教學為目的，從零基礎的預備知識出發，逐層拆解了一套模組化、以 ROS 為基礎的 "
        "TurtleBot3 完全覆蓋路徑規劃系統。其設計亮點可歸納為四點：(1) 以 PCA 自動對齊掃描方向，"
        "使覆蓋對房間朝向具強健性，消除階梯狀鋸齒；(2) 演算法與 ROS 徹底解耦，形成可重用、可測試"
        "的純演算法核心；(3) 透過 move_base 的容錯執行，使單點失敗不致中止任務；(4) 即時多圖層"
        "網頁介面，讓使用者一目了然並一鍵控制。讀完本文，讀者應能理解從光達掃描、SLAM 建圖、"
        "路徑規劃到導航執行與視覺化的完整鏈路，並具備自行修改、擴充與除錯的基礎。未來可朝多區域"
        "分解、路點順序最佳化、frontier 自主探索與線上覆蓋率指標等方向延伸。")

    # ════════════════════════════════════════════════
    #  參考文獻
    # ════════════════════════════════════════════════
    h1(doc, "參考文獻")
    for i, (text, url) in enumerate(REFS, 1):
        ref(doc, i, text, url)

    # ════════════════════════════════════════════════
    #  附錄
    # ════════════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "附錄 A　檔案地圖")
    code(doc,
        "turtlebot3_ccpp_local/\n"
        "├── CMakeLists.txt              catkin 建置設定\n"
        "├── package.xml                 ROS 套件描述與相依\n"
        "├── launch/\n"
        "│   └── start.launch            一鍵啟動全系統\n"
        "├── scripts/\n"
        "│   ├── coverage_planner.py     ★ 純演算法核心（PCA + 牛耕，無 ROS）\n"
        "│   ├── boustrophedon.py        模組 B：規劃節點 → /coverage_path\n"
        "│   └── map_server.py           模組 A：地圖伺服器 + 網頁 + 任務執行\n"
        "├── web/\n"
        "│   └── index.html              前端：Canvas 多圖層即時監控\n"
        "├── test/\n"
        "│   ├── test_lidar_freq.py      光達頻率量測工具\n"
        "│   └── test_lidar_range.py     光達距離範圍量測工具\n"
        "└── docs/                       說明文件（含本論文）")

    h1(doc, "附錄 B　參數總表")
    table(doc,
        ["參數", "所在檔案", "值", "意義"],
        [
            ["SPACING", "coverage_planner.py", "0.18 m", "掃描線間距（留約 10% 重疊）"],
            ["MARGIN", "coverage_planner.py", "0.10 m", "安全邊距（約機器人半徑）"],
            ["CROP_PAD", "map_server.py", "10 格", "地圖裁切邊距"],
            ["REPLAN_INTERVAL", "boustrophedon.py", "5.0 s", "重新規劃最短間隔"],
            ["port", "start.launch / map_server", "8080", "網頁伺服器埠號"],
            ["軌跡記點間距", "map_server.py", "0.1 m", "每移動 0.1 m 記一個軌跡點"],
            ["軌跡上限", "map_server.py", "10000 點", "超過則丟棄最舊點"],
            ["網頁輪詢週期", "web/index.html", "1 s", "每秒抓一次圖與狀態"],
        ],
        widths=[1.7, 2.0, 1.0, 2.0])

    h1(doc, "附錄 C　名詞對照表（中英）")
    table(doc,
        ["中文", "英文 / 原文"],
        [
            ["完全覆蓋路徑規劃", "Complete Coverage Path Planning (CCPP)"],
            ["牛耕式", "Boustrophedon"],
            ["主成分分析", "Principal Component Analysis (PCA)"],
            ["同時定位與建圖", "Simultaneous Localization and Mapping (SLAM)"],
            ["佔據網格", "Occupancy Grid"],
            ["里程計", "Odometry"],
            ["導航堆疊", "Navigation Stack (move_base)"],
            ["動態視窗法", "Dynamic Window Approach (DWA)"],
            ["局部/全域成本地圖", "Local / Global Costmap"],
            ["脫困行為", "Recovery Behaviors"],
            ["形態學（侵蝕/膨脹）", "Morphology (Erosion / Dilation)"],
            ["合作式取消", "Cooperative Cancellation"],
            ["回呼函式", "Callback"],
            ["發布/訂閱", "Publish / Subscribe"],
            ["座標轉換", "Transform (TF)"],
        ],
        widths=[2.6, 4.0])

    out = os.path.join(HERE, "學習論文_CN.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print("生成完成：", path)
