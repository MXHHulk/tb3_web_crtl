# -*- coding: utf-8 -*-
"""
依本專案內容生成一份 A3 直式「專題海報」Word 檔。
版面：頂部彩色標題橫幅 → 全寬「系統架構圖」（分層方塊圖為主、文字為輔）
      → 雙欄核心方法（三法各自成節）與 ROS 概念 → 全寬結果結論 → 參考文獻。
輸出：
  docs/03_海報/專題海報.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.realpath(__file__))
OUT_NAME = os.environ.get("POSTER_OUT", "專題海報.docx")

# ── 配色 ──────────────────────────────────────────────
NAVY   = "1F3A5F"   # 標題橫幅底色（深藍）
BLUE   = "2E6FB0"   # 分區標題底色（藍）
LIGHT  = "EAF1F8"   # 內文區塊淡藍底
ACCENT = "C0392B"   # 重點強調（紅）
GREY   = "5A6672"   # 次要文字（灰）
WHITE  = "FFFFFF"

# 四層角色配色（與系統架構圖對應）
HW_C,  HW_BG = "8A6D3F", "F3ECDD"   # 硬體（棕）
RD_C,  RD_BG = "3B4859", "E4E9EF"   # 現成套件（灰藍）
MY_C,  MY_BG = "5B34B0", "ECE2FB"   # 我寫的程式（紫，全圖唯一亮色）
US_C,  US_BG = "1F6F60", "E1F0EB"   # 使用者端（綠）


# ── 樣式工具 ──────────────────────────────────────────
def set_cjk(run, font_cn="標楷體", font_en="Times New Roman"):
    """同時設定西文與中日韓字型。"""
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)


def shade(el, hex_color):
    """對段落 (w:pPr) 或表格儲存格 (w:tcPr) 設定底色。"""
    pr = el.get_or_add_pPr() if el.tag.endswith('}p') else el._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pr.append(shd)


def cell_shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)


def cell_margins(cell, top=80, bottom=80, left=140, right=140):
    """設定儲存格內距（單位 dxa，1/20 pt）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement(f'w:{name}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def _borders(el, edges, sz=8, color="AAAAAA", val="single"):
    """對段落 (pBdr) 加框線。"""
    pPr = el.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in edges:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '4')
        e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def tbl_borders(table, color, sz=12):
    """整張表格四周加彩色外框（sz 單位為 1/8 pt：12=1.5pt、18=2.25pt）。"""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        if edge in ('insideH', 'insideV'):
            e.set(qn('w:val'), 'none')
        else:
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), str(sz))
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def cell_borders(cell, color, sz=8):
    """單一儲存格四周加彩色框（會覆蓋表格層級的框線設定）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def fix_width(table, widths_cm):
    """關閉自動調整並固定各欄寬度（Word 才不會依內容亂縮放）。"""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)
        for row in table.rows:
            row.cells[i].width = Cm(w)


def left_bar(p, hex_color, size=18):
    """在段落左緣加一條彩色直條（以粗左框線模擬）。"""
    _borders(p._p, ('left',), sz=size, color=hex_color)


def run(p, text, size, color=None, bold=False, cjk=True, font_cn="標楷體"):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if cjk:
        set_cjk(r, font_cn)
    return r


def tight(p, before=0, after=2, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# ── 版面元件 ──────────────────────────────────────────
def banner(doc, title, subtitle, author):
    """頂部整條彩色標題橫幅（單格表格）。"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    c = t.cell(0, 0)
    cell_shade(c, NAVY)
    cell_margins(c, top=110, bottom=110, left=200, right=200)
    c.paragraphs[0].text = ''

    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, title, 25, WHITE, bold=True)
    tight(p, after=4)

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, subtitle, 12.5, "BFD3E6", bold=False)
    tight(p2, after=6)

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p3, author, 11, WHITE, bold=False)
    tight(p3, after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def full_cell(doc, after_gap=3):
    """加入一個全寬、無框線的單格表格，回傳其 cell。"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    t.autofit = False
    c = t.cell(0, 0)
    c.width = Cm(27.3)
    cell_margins(c, left=120, right=120)
    c.paragraphs[0].text = ''
    return c


def section(cell, title, num=None):
    """分區彩色標題列。"""
    p = cell.add_paragraph()
    shade(p._p, BLUE)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(4)
    run(p, title, 13.5, WHITE, bold=True)


def body(cell, text, after=2, size=10):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p, text, size)
    tight(p, after=after, line=1.06)
    return p


def bullet(cell, lead, rest, after=1, lead_color=ACCENT, size=10):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    run(p, "▪ ", size, BLUE, bold=True)
    if lead:
        run(p, lead, size, lead_color, bold=True)
    run(p, rest, size)
    tight(p, after=after, line=1.06)
    return p


def mono_box(cell, lines, size=7.5, after=2):
    """等寬字排版的小示意圖區塊（淡藍底 + 細框）。"""
    p = cell.add_paragraph()
    shade(p._p, LIGHT)
    _borders(p._p, ('top', 'bottom', 'left', 'right'), sz=4, color="C7D6E8")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(after)
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.size = Pt(size)
        r.font.name = "Consolas"
        set_cjk(r, "標楷體", "Consolas")
        if i < len(lines) - 1:
            p.add_run().add_break()


# ── 系統架構圖（真方塊圖：巢狀表格 + 彩色外框） ─────────
#    版面由上而下：① 使用者端 → ② 自行開發 → ③ 現成套件 → ④ 硬體
#    左側紅色 ↓ 為指令下行、右側灰色 ↑ 為感測與狀態上行，構成閉環。
ARCH_W = 24.0     # 層方塊寬（cm）
SUB_W  = 10.8     # ② 層內子方塊寬（cm）
GAP_W  = 2.0      # 子方塊之間的箭頭欄寬（cm）


def _fig_table(container, rows, cols):
    """在 container 內加圖表用表格。

    container 為表格 cell 時，python-docx 會在表格後自動補一段空白段落
    （Word 規定 cell 不能以表格結尾），必須壓成 2 pt，否則每層方塊之間
    會多出一整行高度，海報就會被撐成兩頁。
    """
    t = container.add_table(rows, cols)
    if hasattr(container, '_tc'):
        tail = container.paragraphs[-1]
        pf = tail.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(2)
        tail.add_run().font.size = Pt(1)
    return t


def arch_legend(container, width_cm=ARCH_W):
    """圖例列（靠右，色塊 + 說明）。"""
    t = _fig_table(container, 1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    fix_width(t, [width_cm])
    c = t.cell(0, 0)
    cell_margins(c, top=0, bottom=0, left=0, right=0)
    p = c.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for color, label in ((MY_C, "本專題自行開發"), (RD_C, "ROS 現成套件"),
                         (HW_C, "機器人硬體"), (US_C, "使用者端")):
        run(p, "■ ", 10, color, bold=True)
        run(p, label + "　", 8.5, GREY)
    tight(p, after=1)
    return t


def arch_box(container, tag, note, lines, color, bg,
             width_cm=ARCH_W, border=12, title_pt=12, body_pt=9.5):
    """一層方塊：彩色外框 + 淡色底 + 標題列 + 內容行。"""
    t = _fig_table(container, 1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_borders(t, color, sz=border)
    fix_width(t, [width_cm])
    c = t.cell(0, 0)
    cell_shade(c, bg)
    cell_margins(c, top=24, bottom=24, left=200, right=200)

    p = c.paragraphs[0]
    p.text = ''
    run(p, tag, title_pt, color, bold=True)
    if note:
        run(p, "　　" + note, 9, GREY)
    tight(p, after=2, line=1.0)

    for ln in lines:
        q = c.add_paragraph()
        run(q, ln, body_pt, "222222")
        tight(q, after=0, line=1.1)
    return t


def arch_sub(cell, title, note, lines, color):
    """② 層內的子方塊（白底 + 紫框）。"""
    cell_shade(cell, WHITE)
    cell_borders(cell, color, sz=8)
    cell_margins(cell, top=35, bottom=35, left=140, right=140)
    p = cell.paragraphs[0]
    p.text = ''
    run(p, title, 11, color, bold=True)
    tight(p, after=0, line=1.0)
    if note:
        q = cell.add_paragraph()
        run(q, note, 9, GREY)
        tight(q, after=2, line=1.0)
    for ln in lines:
        q = cell.add_paragraph()
        run(q, ln, 9, "222222")
        tight(q, after=0, line=1.08)


def arch_core(container, tag, note, left, right, mid_lbl, color, bg,
              width_cm=ARCH_W):
    """② 本專題自行開發：外框方塊內含兩個子方塊與水平箭頭。"""
    t = _fig_table(container, 1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_borders(t, color, sz=18)                     # 2.25 pt，全圖最粗
    fix_width(t, [width_cm])
    c = t.cell(0, 0)
    cell_shade(c, bg)
    cell_margins(c, top=60, bottom=40, left=140, right=140)

    p = c.paragraphs[0]
    p.text = ''
    run(p, tag, 12, color, bold=True)
    run(p, "　　" + note, 9, GREY)
    tight(p, after=2, line=1.0)

    inner = c.add_table(1, 3)
    no_borders(inner)
    inner.alignment = WD_TABLE_ALIGNMENT.CENTER
    fix_width(inner, [SUB_W, GAP_W, SUB_W])
    arch_sub(inner.cell(0, 0), *left, color=color)
    arch_sub(inner.cell(0, 2), *right, color=color)

    mid = inner.cell(0, 1)
    cell_vcenter(mid)
    cell_margins(mid, top=0, bottom=0, left=0, right=0)
    mp = mid.paragraphs[0]
    mp.text = ''
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(mp, "▶", 16, color, bold=True)
    tight(mp, after=0, line=1.0)
    ml = mid.add_paragraph()
    ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, seg in enumerate(mid_lbl.split("\n")):
        if i:
            ml.add_run().add_break()
        run(ml, seg, 8.5, GREY)
    tight(ml, after=0, line=1.0)

    # 巢狀表格後 python-docx 會補一段空白，壓到最小以免撐高方塊
    tail = c.paragraphs[-1]
    tail.text = ''
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing = 1.0
    for r in tail.runs or [tail.add_run()]:
        r.font.size = Pt(2)
    return t


def arch_flow(container, down_lbl, up_lbl, width_cm=ARCH_W):
    """層與層之間的資料流列：左紅 ↓ 指令、右灰 ↑ 資料。"""
    t = _fig_table(container, 1, 2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    fix_width(t, [width_cm / 2, width_cm / 2])
    l, r = t.cell(0, 0), t.cell(0, 1)
    for c in (l, r):
        cell_margins(c, top=10, bottom=10, left=60, right=60)

    lp = l.paragraphs[0]
    lp.text = ''
    run(lp, "↓　" + down_lbl, 8.5, ACCENT, bold=True)
    tight(lp, before=0, after=0, line=1.0)

    rp = r.paragraphs[0]
    rp.text = ''
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(rp, up_lbl + "　↑", 8.5, GREY, bold=True)
    tight(rp, before=0, after=0, line=1.0)
    return t


def arch_caption(container, text, width_cm=ARCH_W):
    """圖說（置中）。"""
    t = _fig_table(container, 1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    fix_width(t, [width_cm])
    c = t.cell(0, 0)
    cell_margins(c, top=0, bottom=0, left=0, right=0)
    p = c.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, text, 10, "222222", bold=True)
    for color, label in ((MY_C, "本專題自行開發"), (RD_C, "ROS 現成套件"),
                         (HW_C, "機器人硬體"), (US_C, "使用者端")):
        run(p, "　■ ", 9, color, bold=True)
        run(p, label, 8, GREY)
    tight(p, before=3, after=0, line=1.0)
    return t


def build_arch_figure(container):
    """把整張系統架構圖畫進 container（Document 或表格 cell）。"""

    arch_box(container, "① 使用者端", "免接螢幕、免安裝軟體", [
        "瀏覽器（手機／電腦）：▶ 開始覆蓋 / ■ 停止・即時進度（第 n / N 點）　│　7 圖層可各自開關："
        "原始地圖・侵蝕・膨脹・覆蓋路徑・軌跡・雷射・機器人",
    ], US_C, US_BG)

    arch_flow(container,
              "HTTP POST　/coverage/start、/coverage/stop",
              "HTTP GET　地圖 PNG、/robot_state、/scan（JSON）")

    arch_core(container, "② 本專題自行開發", "核心貢獻",
              left=("coverage_planner.py", "純演算法・不依賴 ROS", [
                  "1　安全邊距膨脹　MARGIN = 0.10 m",
                  "2　PCA 主軸對齊　長軸掃描／短軸換行",
                  "3　牛耕路點生成　SPACING = 0.18 m（10% 重疊）",
              ]),
              right=("map_server.py（ROS 節點）", "系統樞紐", [
                  "地圖影像處理：侵蝕／膨脹／自動裁切",
                  "逐點送導航目標・失敗跳過容錯・可隨時中止",
                  "Flask 網頁伺服器 :8080　│　threading.Lock 保護共用資料",
              ]),
              mid_lbl="路點串列\n[(x, y), …]",
              color=MY_C, bg=MY_BG)

    arch_flow(container,
              "move_base action：逐點目標 goal（含朝向）",
              "/map　·　/tf (map→odom)　·　/scan　·　/odom")

    arch_box(container, "③ ROS 現成套件", "經驗證的生態系・非本專題開發", [
        "gmapping　SLAM 建圖＋定位　│　move_base + DWA　全域規劃＋即時避障　│　turtlebot3_bringup　硬體驅動　│　tf2　座標轉換樹",
    ], RD_C, RD_BG)

    arch_flow(container, "/cmd_vel　輪速指令", "雷射距離、輪子里程（USB 序列）")

    arch_box(container, "④ 機器人硬體", "TurtleBot3 Burger", [
        "LDS-01 360° 雷射雷達　│　輪子編碼器　│　OpenCR + Dynamixel 馬達",
    ], HW_C, HW_BG)

    arch_caption(container, "圖 1　系統四層架構與資料流（左：指令下行；右：感測與狀態上行）")



# ── 建立海報 ──────────────────────────────────────────
def build():
    doc = Document()
    # A3 直式 + 窄邊界
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(42.0)
    sec.top_margin = Cm(1.1)
    sec.bottom_margin = Cm(1.1)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)

    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.makeelement(qn('w:rFonts'), {})
    rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), "標楷體")

    # ═══ 標題橫幅 ═══
    banner(doc,
           "TurtleBot3 牛耕式完全覆蓋路徑規劃（CCPP）系統",
           "基於 ROS 與 PCA 主軸對齊之牛耕式演算法，及其即時網頁視覺化",
           "作者：MXHHulk　│　機器人與自動化　│　hulk71777@gmail.com　│　ROS Noetic · TurtleBot3 Burger")

    # ═══ 全寬：研究動機 ═══
    top = full_cell(doc)
    section(top, "一、研究動機與目的")
    body(top, "完全覆蓋路徑規劃（CCPP）要求移動機器人走遍工作空間中每一個可到達區域，是掃地機器人、"
              "地面巡檢與農業自動化的核心能力。牛耕式（來回掃描）最大的難題是「掃描線該朝哪個方向」："
              "地圖座標系由開機朝向決定，牆面常與地圖軸向傾斜，沿固定軸向掃會切出階梯狀鋸齒、大量轉彎。"
              "本專題打造一套可部署、模組化的 CCPP 系統，整合 SLAM → 規劃 → 執行 → 網頁視覺化的完整"
              "流程，並以 PCA 自動對齊房間長軸，適應任意朝向的空間。", after=2)

    # ═══ 全寬：系統架構（分層方塊圖為主，文字為輔） ═══
    arch = full_cell(doc)
    section(arch, "二、系統架構")
    body(arch, "系統由四層協作構成：指令由上而下（左側紅色箭頭），感測資料與執行狀態由下而上"
               "（右側灰色箭頭），形成閉環。圖中紫色方塊為本專題自行開發，其餘為硬體與 ROS 現成套件。",
         after=3, size=10)

    build_arch_figure(arch)

    bullet(arch, "互動關鍵：", "現成套件只會「從 A 走到 B」，不會回答「怎麼走才能不重複走遍整個房間」——"
                            "全覆蓋路線由 ② 層的 coverage_planner 算出，逐點排隊、監控、容錯與中止亦由 "
                            "map_server 實作；coverage_planner 完全不 import rospy，可獨立測試並被多節點共用。",
           lead_color=MY_C, after=2)

    # ═══ 雙欄：核心方法（三法拆開）＋ ROS ═══
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tbl)
    tbl.autofit = False
    L, R = tbl.cell(0, 0), tbl.cell(0, 1)
    for c in (L, R):
        c.width = Cm(13.4)
        cell_margins(c, left=120, right=120)
    L.paragraphs[0].text = ''
    R.paragraphs[0].text = ''

    # ── 左欄 ──
    section(L, "三、核心方法①　PCA 主軸對齊")
    body(L, "以主成分分析（PCA）自「可走格點集」的共變異數矩陣求特徵向量，自動找出自由空間的"
            "長軸與短軸：")
    bullet(L, "長軸 → 掃描方向：", "沿空間最長方向來回走，掃描線最少、轉彎最少。")
    bullet(L, "短軸 → 換行方向：", "每掃完一條線往短軸平移一個間距。")
    body(L, "因軸向源自空間形狀，不論四邊形相對地圖是正是斜，掃描線都自動貼齊長邊——這就是能"
            "「走遍任意方向四邊形」的關鍵。")
    mono_box(L, [
        "  固定水平掃(笨)        PCA 對齊掃(本專案)",
        "  ┌──────────┐          ┌──────────┐",
        "  │ ─ ─ ─ ─  │ 斜房間   │ ╲ ╲ ╲ ╲  │ 貼齊長軸",
        "  │─ ─ ─ ─   │ 切碎、   │  ╲ ╲ ╲ ╲ │ 完整又順",
        "  │ ─ ─ ─ ─  │ 漏邊角   │ ╲ ╲ ╲ ╲  │",
        "  └──────────┘          └──────────┘",
    ])

    section(L, "四、核心方法②　安全邊距膨脹")
    body(L, "機器人有體積（直徑約 0.20 m），路徑不能貼牆。規劃前先把障礙物「膨脹」約一個機器人"
            "半徑，讓路徑對「點機器人」安全＝實體機器人遠離牆面。")
    mono_box(L, [
        "obs  = data == 100                 # 障礙格",
        "r    = round(margin / resolution)  # 0.10/0.05 ≈ 2 格",
        "kern = np.ones((2*r+1, 2*r+1))",
        "safe = binary_dilation(obs, kern)  # 障礙向外長胖",
        "free = (data == 0) & ~safe         # 真正可走區",
    ])
    bullet(L, "實作：", "scripts/coverage_planner.py 的 apply_safety_margin()，MARGIN=0.10 m。")
    bullet(L, "注意：", "網頁的侵蝕／膨脹圖層是固定核、純視覺參考，與規劃用的安全邊距已解耦，勿混淆。")

    # ── 右欄 ──
    section(R, "五、核心方法③　牛耕路點生成")
    body(R, "把每個可走點投影到長／短軸，沿短軸每隔一個間距切一條「帶」，取該帶在長軸上投影的"
            "最小／最大值作為線段兩端，方向逐條交替，串成連續蛇行路線。")
    mono_box(R, [
        "  短軸每隔 SPACING 切一條帶，長軸取端點",
        "  ┌────────────────────┐",
        "  │ ►►►►►►►►►►►►►►►►►►► │ 第1條 沿長軸",
        "  │ ◄◄◄◄◄◄◄◄◄◄◄◄◄◄◄◄◄◄◄ │ 第2條 反向",
        "  │ ►►►►►►►►►►►►►►►►►►► │ 交替蛇行到底",
        "  └────────────────────┘",
    ])
    bullet(R, "10% 重疊：", "SPACING=0.18 m 略小於覆蓋寬度 0.20 m，相鄰掃描線不留縫，確保全覆蓋。")
    bullet(R, "算得快：", "全程 numpy 向量化（np.where、投影、遮罩），無逐格迴圈，上萬格也能瞬間"
                       "重算——收到新地圖即時重規劃的基礎。")

    section(R, "六、支撐本專題的 ROS 概念")
    body(R, "ROS（機器人作業系統）不是一般作業系統，而是讓「許多小程式協同運作」的框架——正因如此，"
            "上面的現成套件與我寫的程式才能無縫接在一起。")
    bullet(R, "節點 Node：", "每個功能是一支獨立程式；gmapping、move_base、map_server 同時運作。")
    bullet(R, "話題 Topic：", "節點以「發布／訂閱」非同步傳資料（/scan /map /odom），我的節點訂閱取用。")
    bullet(R, "動作 Action：", "適合「需時間、要回報、可取消」的任務；我用 move_base action 逐點下達導航。")
    bullet(R, "座標轉換 TF：", "維護 map／odom／base_link 相對位置，隨時可查機器人在地圖哪裡。")
    bullet(R, "launch：", "一鍵啟動所有節點與參數：roslaunch turtlebot3_ccpp start.launch。")

    # ═══ 全寬：結果與結論 ═══
    res = full_cell(doc)
    section(res, "七、結果與結論")
    bullet(res, "PCA 對齊：", "實體 TurtleBot3 Burger 上，掃描線穩定與牆面平行，即使地圖旋轉亦然，"
                          "確實消除階梯鋸齒、減少轉彎。")
    bullet(res, "安全與容錯：", "安全邊距使軌跡遠離牆面；逐點容錯讓個別目標暫時到不了仍能完成整體覆蓋。")
    bullet(res, "模組化價值：", "演算法與 ROS 解耦，覆蓋邏輯可獨立演練調校，換 UI 不必改演算法。")
    bullet(res, "未來工作：", "非凸環境多區域分解、路點順序 TSP 最佳化、frontier 自主探索、"
                          "網頁即時覆蓋率指標。")

    # ═══ 底部參考文獻列 ═══
    ref = doc.add_paragraph()
    shade(ref._p, LIGHT)
    ref.paragraph_format.space_before = Pt(2)
    run(ref, "參考文獻：", 8.5, BLUE, bold=True)
    run(ref, "[1] Galceran & Carreras, CPP survey, RAS 2013. "
             "[2] Choset & Pignon, Boustrophedon decomposition, FSR 1998. "
             "[3] Grisetti et al., gmapping, T-RO 2007. "
             "[4] Fox et al., Dynamic Window Approach, RAM 1997. "
             "[5] Quigley et al., ROS, ICRA 2009. "
             "[7] Jolliffe & Cadima, PCA review, Phil. Trans. R. Soc. A 2016.",
        8.5, GREY)
    tight(ref, after=0, line=1.05)

    out = os.path.join(HERE, OUT_NAME)
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print("生成完成：", path)
